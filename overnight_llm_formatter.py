#!/usr/bin/env python3
"""
Overnight LLM Document Formatter
Designed for unattended batch processing with your local Ollama setup
"""

import requests
import json
import time
import os
import pickle
from pathlib import Path
from docx import Document
from collections import Counter
from datetime import datetime
import signal
import sys

class OvernightLLMFormatter:
    """LLM-based document formatter optimized for overnight processing."""
    
    def __init__(self, ollama_host="http://localhost:11434", model="gemma3:27b-it-qat"):
        self.ollama_host = ollama_host
        self.model = model
        self.progress_file = "overnight_progress.pkl"
        self.log_file = "overnight_processing.log"
        self.processed_paragraphs = []
        self.current_index = 0
        self.start_time = None
        self.llm_calls = 0
        self.rule_hits = 0
        
        # Setup signal handler for graceful shutdown
        signal.signal(signal.SIGINT, self._signal_handler)
        signal.signal(signal.SIGTERM, self._signal_handler)
        
        # Style classification prompts
        self.classification_prompt = """You are a document formatting expert. Classify this paragraph into one of these exact styles:

STYLES:
- Heading 1: Major section headers (EMPLOYMENT POLICIES, SCHEDULING, etc.)
- Heading 2: Section headers with (RC-THPPH) codes, policy statements
- Heading 3: Subsection headers, requirements, procedures
- Heading 4: Details, questions, notes
- Heading 5: Options, system details, minor classifications
- Body Text: Explanatory content, descriptions, policy explanations
- List Paragraph: Instructions, requirements, action items, procedures
- Normal: Special formatting, quotes, references

PARAGRAPH TO CLASSIFY:
"{text}"

CONTEXT:
- This is from a policy handbook/manual
- Previous paragraph: "{previous_text}"
- Document section: Policy and procedures

INSTRUCTIONS:
- Consider the content meaning and intent
- Look for imperative language (must, shall, should) = often List Paragraph
- Look for ALL CAPS = usually headings
- Look for explanatory content = usually Body Text
- Respond with ONLY the exact style name

STYLE:"""

        self.filtering_prompt = """You are a document content expert. Determine if this text should be INCLUDED or FILTERED from the final document.

FILTER OUT (remove):
- Headers and footers
- Page numbers
- "INTENTIONALLY LEFT BLANK" pages
- Table of contents entries
- Navigation elements
- Revision information
- Company headers
- Document titles in headers

KEEP (include):
- Policy content
- Procedures
- Requirements
- Explanations
- Lists and instructions
- Actual document content

TEXT TO EVALUATE:
"{text}"

INSTRUCTIONS:
- Consider the content meaning and purpose
- Navigation/metadata should be filtered
- Actual policy content should be kept
- Respond with ONLY: INCLUDE or FILTER

DECISION:"""

    def _signal_handler(self, signum, frame):
        """Handle graceful shutdown."""
        self.log(f"Received signal {signum}, saving progress...")
        self._save_progress()
        sys.exit(0)
    
    def log(self, message):
        """Log message with timestamp."""
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        log_entry = f"[{timestamp}] {message}"
        print(log_entry)
        
        with open(self.log_file, "a") as f:
            f.write(log_entry + "\n")
    
    def _save_progress(self):
        """Save current progress to file."""
        progress_data = {
            'processed_paragraphs': self.processed_paragraphs,
            'current_index': self.current_index,
            'llm_calls': self.llm_calls,
            'rule_hits': self.rule_hits,
            'start_time': self.start_time
        }
        
        with open(self.progress_file, 'wb') as f:
            pickle.dump(progress_data, f)
        
        self.log(f"Progress saved: {self.current_index} paragraphs processed")
    
    def _load_progress(self):
        """Load progress from file if it exists."""
        if os.path.exists(self.progress_file):
            try:
                with open(self.progress_file, 'rb') as f:
                    progress_data = pickle.load(f)
                
                self.processed_paragraphs = progress_data.get('processed_paragraphs', [])
                self.current_index = progress_data.get('current_index', 0)
                self.llm_calls = progress_data.get('llm_calls', 0)
                self.rule_hits = progress_data.get('rule_hits', 0)
                self.start_time = progress_data.get('start_time', None)
                
                self.log(f"Resumed from progress: {self.current_index} paragraphs already processed")
                return True
            except Exception as e:
                self.log(f"Error loading progress: {e}")
                return False
        return False
    
    def test_ollama_connection(self):
        """Test connection to Ollama server."""
        try:
            response = requests.get(f"{self.ollama_host}/api/tags")
            if response.status_code == 200:
                models = response.json().get('models', [])
                model_names = [model['name'] for model in models]
                
                if self.model in model_names:
                    self.log(f"✅ Connected to Ollama with model {self.model}")
                    return True
                else:
                    self.log(f"❌ Model {self.model} not found. Available: {model_names}")
                    return False
            else:
                self.log(f"❌ Ollama server not responding: {response.status_code}")
                return False
        except Exception as e:
            self.log(f"❌ Connection error: {e}")
            return False
    
    def call_ollama(self, prompt, max_retries=3, timeout=30):
        """Call Ollama with retry logic and timeout."""
        for attempt in range(max_retries):
            try:
                payload = {
                    "model": self.model,
                    "prompt": prompt,
                    "stream": False,
                    "options": {
                        "temperature": 0.1,  # Low temperature for consistency
                        "top_k": 10,
                        "top_p": 0.3
                    }
                }
                
                response = requests.post(
                    f"{self.ollama_host}/api/generate",
                    json=payload,
                    timeout=timeout
                )
                
                if response.status_code == 200:
                    result = response.json()
                    self.llm_calls += 1
                    return result.get('response', '').strip()
                else:
                    self.log(f"Ollama error (attempt {attempt + 1}): {response.status_code}")
                    
            except requests.exceptions.Timeout:
                self.log(f"Timeout on attempt {attempt + 1}")
            except Exception as e:
                self.log(f"Error on attempt {attempt + 1}: {e}")
            
            if attempt < max_retries - 1:
                time.sleep(2 ** attempt)  # Exponential backoff
        
        return None
    
    def quick_rule_filter(self, text):
        """Quick rule-based filtering for obvious cases."""
        text_lower = text.lower().strip()
        
        # Obvious filters
        if len(text_lower) < 3:
            return False
        
        # Critical patterns
        if 'intentionally left blank' in text_lower:
            return False
        
        if any(pattern in text_lower for pattern in [
            'table of contents', 'page ', 'revision:', 'report #',
            'company name', 'oae.', 'master table'
        ]):
            return False
        
        # Navigation patterns
        if text_lower.count('.') > 5:  # Dotted lines
            return False
        
        return True
    
    def quick_rule_classify(self, text):
        """Quick rule-based classification for obvious cases."""
        text_clean = text.strip()
        text_lower = text_clean.lower()
        
        # High confidence patterns
        
        # List paragraphs - strong indicators
        if any(text_clean.startswith(prefix) for prefix in ['•', '-', '*', '◦', '►']):
            return "List Paragraph", 0.95
        
        if text_clean.startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.')):
            return "List Paragraph", 0.95
        
        # Imperative verbs
        if any(text_lower.startswith(verb) for verb in [
            'read,', 'understand', 'comply', 'report', 'maintain', 'avoid',
            'use', 'respect', 'promote', 'cooperate', 'ensure', 'follow'
        ]):
            return "List Paragraph", 0.90
        
        # Modal verbs
        if any(phrase in text_lower for phrase in ['must ', 'shall ', 'will ', 'should ']):
            return "List Paragraph", 0.85
        
        # Headings - strong indicators
        if text_clean.isupper() and len(text_clean) < 100 and not text_clean.endswith('.'):
            if len(text_clean) < 30:
                return "Heading 2", 0.90
            else:
                return "Heading 1", 0.90
        
        # Appendix pattern
        if text_clean.startswith('Appendix '):
            return "Heading 1", 0.95
        
        # Special patterns
        if text_clean.startswith('NOTE:'):
            return "Heading 4", 0.90
        
        if text_clean.endswith('?'):
            return "Heading 4", 0.85
        
        return None, 0.0
    
    def llm_filter(self, text):
        """Use LLM to determine if content should be filtered."""
        prompt = self.filtering_prompt.format(text=text[:500])  # Truncate for efficiency
        
        response = self.call_ollama(prompt, timeout=15)
        if response:
            if 'FILTER' in response.upper():
                return False
            elif 'INCLUDE' in response.upper():
                return True
        
        # Default to include on LLM failure
        return True
    
    def llm_classify(self, text, previous_text=""):
        """Use LLM to classify paragraph style."""
        prompt = self.classification_prompt.format(
            text=text[:800],  # Truncate for efficiency
            previous_text=previous_text[:200] if previous_text else ""
        )
        
        response = self.call_ollama(prompt, timeout=20)
        if response:
            # Extract style from response
            response_upper = response.upper()
            styles = [
                "HEADING 1", "HEADING 2", "HEADING 3", "HEADING 4", "HEADING 5",
                "BODY TEXT", "LIST PARAGRAPH", "NORMAL"
            ]
            
            for style in styles:
                if style in response_upper:
                    return style.title()
        
        # Default fallback
        return "Body Text"
    
    def process_document(self, input_path, output_path, resume=True):
        """Process document with overnight LLM processing."""
        self.log("🌙 Starting overnight LLM document processing...")
        
        # Test Ollama connection
        if not self.test_ollama_connection():
            self.log("❌ Cannot connect to Ollama. Exiting.")
            return False
        
        # Load progress if resuming
        if resume and self._load_progress():
            self.log("📄 Resuming from previous session...")
        else:
            self.log("📄 Starting fresh processing...")
            self.start_time = time.time()
        
        # Load document
        doc = Document(input_path)
        all_paragraphs = [p for p in doc.paragraphs if p.text.strip()]
        
        self.log(f"📊 Document loaded: {len(all_paragraphs)} total paragraphs")
        
        # Process paragraphs
        for i, para in enumerate(all_paragraphs):
            if i < self.current_index:
                continue  # Skip already processed
            
            text = para.text.strip()
            previous_text = all_paragraphs[i-1].text.strip() if i > 0 else ""
            
            # Progress logging
            if i % 10 == 0:
                elapsed = time.time() - self.start_time if self.start_time else 0
                rate = self.llm_calls / (elapsed / 60) if elapsed > 0 else 0
                self.log(f"📈 Progress: {i}/{len(all_paragraphs)} ({(i/len(all_paragraphs)*100):.1f}%) - LLM calls: {self.llm_calls} - Rate: {rate:.1f}/min")
            
            # Step 1: Quick rule-based filtering
            if not self.quick_rule_filter(text):
                self.rule_hits += 1
                continue  # Skip filtered content
            
            # Step 2: Quick rule-based classification
            rule_style, confidence = self.quick_rule_classify(text)
            
            if rule_style and confidence > 0.85:
                # Use rule-based result
                style = rule_style
                self.rule_hits += 1
            else:
                # Step 3: LLM filtering for complex cases
                if not self.llm_filter(text):
                    continue  # Skip LLM-filtered content
                
                # Step 4: LLM classification
                style = self.llm_classify(text, previous_text)
            
            # Store result
            self.processed_paragraphs.append({
                'text': text,
                'style': style,
                'index': i
            })
            
            self.current_index = i + 1
            
            # Save progress every 50 paragraphs
            if i % 50 == 0:
                self._save_progress()
        
        # Create output document
        self.log("📝 Creating formatted document...")
        new_doc = Document()
        style_counts = Counter()
        
        for item in self.processed_paragraphs:
            new_para = new_doc.add_paragraph(item['text'])
            
            try:
                new_para.style = item['style']
                style_counts[item['style']] += 1
            except:
                new_para.style = 'Body Text'
                style_counts['Body Text'] += 1
        
        # Save document
        new_doc.save(output_path)
        
        # Final statistics
        total_time = time.time() - self.start_time if self.start_time else 0
        self.log(f"✅ Processing complete!")
        self.log(f"📊 Final stats:")
        self.log(f"   • Total paragraphs processed: {len(self.processed_paragraphs)}")
        self.log(f"   • Total paragraphs filtered: {len(all_paragraphs) - len(self.processed_paragraphs)}")
        self.log(f"   • LLM calls made: {self.llm_calls}")
        self.log(f"   • Rule hits: {self.rule_hits}")
        self.log(f"   • Processing time: {total_time/3600:.1f} hours")
        self.log(f"   • Average rate: {self.llm_calls/(total_time/60):.1f} LLM calls/minute")
        
        # Style distribution
        self.log(f"📈 Style distribution:")
        for style, count in style_counts.most_common():
            percentage = (count / len(self.processed_paragraphs)) * 100
            self.log(f"   • {style}: {count} ({percentage:.1f}%)")
        
        # Clean up progress file
        if os.path.exists(self.progress_file):
            os.remove(self.progress_file)
        
        return True

def main():
    """Main function for overnight processing."""
    # File paths
    input_path = "/Users/vincent/Desktop/watson/PPH_original.docx"
    output_path = "/Users/vincent/Desktop/watson/PPH_claude_overnight_llm_formatted.docx"
    
    # Create formatter
    formatter = OvernightLLMFormatter()
    
    # Process document
    success = formatter.process_document(input_path, output_path, resume=True)
    
    if success:
        print("\n🎉 Overnight processing completed successfully!")
        print(f"📄 Output saved to: {output_path}")
        print(f"📋 Processing log: {formatter.log_file}")
    else:
        print("\n❌ Processing failed. Check logs for details.")

if __name__ == '__main__':
    main()