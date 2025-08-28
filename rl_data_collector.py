#!/usr/bin/env python3
"""
RL Data Collector for Document Formatting System
Captures team actions as State-Action-Reward trajectories for reinforcement learning training
"""

import json
import pickle
import time
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Any, Optional, Tuple
from dataclasses import dataclass, asdict
from docx import Document
from collections import Counter
import hashlib

@dataclass
class DocumentState:
    """Represents the current state of the document being formatted"""
    paragraph_count: int
    style_distribution: Dict[str, int]
    filtered_content_count: int
    document_complexity: float
    current_stage: str
    formatting_rules_applied: List[str]
    user_preferences: Dict[str, Any]
    document_type: str
    timestamp: str

@dataclass
class FormattingAction:
    """Represents a formatting action taken by the system or user"""
    action_type: str  # 'filter', 'classify', 'style_change', 'manual_correction'
    parameters: Dict[str, Any]
    confidence: float
    method_used: str  # 'rule_based', 'llm', 'pattern_match', 'manual'
    paragraph_index: int
    original_text: str
    applied_style: str
    timestamp: str

@dataclass
class Reward:
    """Represents the reward/quality score for the action"""
    quality_score: float  # 0.0 to 1.0
    user_satisfaction: Optional[float]
    compliance_score: float
    efficiency_metric: float
    error_reduction: float
    manual_corrections_needed: int
    timestamp: str

@dataclass
class Episode:
    """Complete formatting session from start to finish"""
    episode_id: str
    user_id: str
    session_start: str
    session_end: Optional[str]
    document_id: str
    input_document_path: str
    output_document_path: str
    method_used: str
    states: List[DocumentState]
    actions: List[FormattingAction]
    rewards: List[Reward]
    metadata: Dict[str, Any]
    total_processing_time: float
    final_quality_score: Optional[float]

class RLDataCollector:
    """Collects team actions and system interactions for RL training data"""
    
    def __init__(self, data_dir: str = "rl_training_data"):
        self.data_dir = Path(data_dir)
        self.data_dir.mkdir(exist_ok=True)
        
        self.current_episode: Optional[Episode] = None
        self.session_start_time: Optional[float] = None
        
        # Data storage files
        self.episodes_file = self.data_dir / "episodes.jsonl"
        self.states_file = self.data_dir / "states.jsonl"
        self.actions_file = self.data_dir / "actions.jsonl"
        self.rewards_file = self.data_dir / "rewards.jsonl"
        self.metadata_file = self.data_dir / "metadata.json"
        
        # Quality metrics tracking
        self.quality_metrics = {
            "total_episodes": 0,
            "total_actions": 0,
            "average_quality_score": 0.0,
            "method_performance": {},
            "user_satisfaction_history": []
        }

    def start_episode(self, user_id: str, document_path: str, method_used: str, 
                     user_preferences: Dict[str, Any] = None) -> str:
        """Start a new formatting episode"""
        episode_id = self._generate_episode_id(user_id, document_path)
        self.session_start_time = time.time()
        
        # Extract initial document state
        initial_state = self._capture_document_state(
            document_path, "initialization", user_preferences or {}
        )
        
        self.current_episode = Episode(
            episode_id=episode_id,
            user_id=user_id,
            session_start=datetime.now().isoformat(),
            session_end=None,
            document_id=self._get_document_id(document_path),
            input_document_path=document_path,
            output_document_path="",
            method_used=method_used,
            states=[initial_state],
            actions=[],
            rewards=[],
            metadata={
                "user_preferences": user_preferences or {},
                "system_version": "1.0.0",
                "initial_document_size": Path(document_path).stat().st_size if Path(document_path).exists() else 0
            },
            total_processing_time=0.0,
            final_quality_score=None
        )
        
        self._log_event("episode_started", {"episode_id": episode_id, "method": method_used})
        return episode_id

    def record_action(self, action_type: str, parameters: Dict[str, Any], 
                     confidence: float, method_used: str, paragraph_index: int = -1,
                     original_text: str = "", applied_style: str = "") -> None:
        """Record a formatting action"""
        if not self.current_episode:
            raise ValueError("No active episode. Call start_episode() first.")
        
        action = FormattingAction(
            action_type=action_type,
            parameters=parameters,
            confidence=confidence,
            method_used=method_used,
            paragraph_index=paragraph_index,
            original_text=original_text[:200],  # Truncate for storage
            applied_style=applied_style,
            timestamp=datetime.now().isoformat()
        )
        
        self.current_episode.actions.append(action)
        self._save_action(action)
        
        # Capture state after action
        if self.current_episode.input_document_path:
            new_state = self._capture_document_state(
                self.current_episode.input_document_path, 
                f"after_{action_type}",
                self.current_episode.metadata.get("user_preferences", {})
            )
            self.current_episode.states.append(new_state)
            self._save_state(new_state)

    def record_reward(self, quality_score: float, compliance_score: float = 0.0,
                     efficiency_metric: float = 0.0, error_reduction: float = 0.0,
                     manual_corrections_needed: int = 0,
                     user_satisfaction: Optional[float] = None) -> None:
        """Record reward for the last action"""
        if not self.current_episode or not self.current_episode.actions:
            raise ValueError("No action to assign reward to.")
        
        reward = Reward(
            quality_score=quality_score,
            user_satisfaction=user_satisfaction,
            compliance_score=compliance_score,
            efficiency_metric=efficiency_metric,
            error_reduction=error_reduction,
            manual_corrections_needed=manual_corrections_needed,
            timestamp=datetime.now().isoformat()
        )
        
        self.current_episode.rewards.append(reward)
        self._save_reward(reward)

    def end_episode(self, output_document_path: str, 
                   final_quality_score: Optional[float] = None) -> Dict[str, Any]:
        """End the current episode and calculate final metrics"""
        if not self.current_episode:
            raise ValueError("No active episode to end.")
        
        self.current_episode.session_end = datetime.now().isoformat()
        self.current_episode.output_document_path = output_document_path
        self.current_episode.total_processing_time = time.time() - (self.session_start_time or 0)
        self.current_episode.final_quality_score = final_quality_score
        
        # Calculate episode summary
        episode_summary = self._calculate_episode_summary()
        
        # Save complete episode
        self._save_episode(self.current_episode)
        
        # Update quality metrics
        self._update_quality_metrics(self.current_episode)
        
        # Reset for next episode
        episode_data = asdict(self.current_episode)
        self.current_episode = None
        self.session_start_time = None
        
        return episode_summary

    def record_user_feedback(self, satisfaction_score: float, 
                           feedback_text: str = "") -> None:
        """Record user feedback for the current episode"""
        if not self.current_episode:
            return
        
        # Add feedback to last reward if exists, otherwise create new reward
        if self.current_episode.rewards:
            self.current_episode.rewards[-1].user_satisfaction = satisfaction_score
        else:
            self.record_reward(
                quality_score=satisfaction_score,
                user_satisfaction=satisfaction_score
            )
        
        self.current_episode.metadata["user_feedback"] = {
            "satisfaction_score": satisfaction_score,
            "feedback_text": feedback_text,
            "timestamp": datetime.now().isoformat()
        }

    def get_training_data(self, format_type: str = "trajectory") -> List[Dict[str, Any]]:
        """Export collected data in various formats for RL training"""
        if format_type == "trajectory":
            return self._export_trajectory_format()
        elif format_type == "state_action_pairs":
            return self._export_state_action_format()
        elif format_type == "episodes":
            return self._export_episodes_format()
        else:
            raise ValueError(f"Unknown format_type: {format_type}")

    def _capture_document_state(self, document_path: str, stage: str, 
                               user_preferences: Dict[str, Any]) -> DocumentState:
        """Capture the current state of the document"""
        try:
            if not Path(document_path).exists():
                # Return default state if document doesn't exist yet
                return DocumentState(
                    paragraph_count=0,
                    style_distribution={},
                    filtered_content_count=0,
                    document_complexity=0.0,
                    current_stage=stage,
                    formatting_rules_applied=[],
                    user_preferences=user_preferences,
                    document_type="unknown",
                    timestamp=datetime.now().isoformat()
                )
            
            doc = Document(document_path)
            paragraphs = [p for p in doc.paragraphs if p.text.strip()]
            
            # Style distribution
            style_dist = Counter()
            for para in paragraphs:
                style_name = para.style.name if para.style else "Normal"
                style_dist[style_name] += 1
            
            # Document complexity (simple heuristic)
            complexity = self._calculate_document_complexity(paragraphs)
            
            return DocumentState(
                paragraph_count=len(paragraphs),
                style_distribution=dict(style_dist),
                filtered_content_count=0,  # Would be calculated by formatter
                document_complexity=complexity,
                current_stage=stage,
                formatting_rules_applied=[],  # Would be populated by formatter
                user_preferences=user_preferences,
                document_type=self._detect_document_type(paragraphs),
                timestamp=datetime.now().isoformat()
            )
        except Exception as e:
            print(f"Error capturing document state: {e}")
            return DocumentState(
                paragraph_count=0,
                style_distribution={},
                filtered_content_count=0,
                document_complexity=0.0,
                current_stage=stage,
                formatting_rules_applied=[],
                user_preferences=user_preferences,
                document_type="error",
                timestamp=datetime.now().isoformat()
            )

    def _calculate_document_complexity(self, paragraphs: List) -> float:
        """Calculate a complexity score for the document"""
        if not paragraphs:
            return 0.0
        
        # Simple complexity metrics
        total_length = sum(len(p.text) for p in paragraphs)
        avg_length = total_length / len(paragraphs)
        unique_styles = len(set(p.style.name if p.style else "Normal" for p in paragraphs))
        
        # Normalize and combine
        complexity = min(1.0, (avg_length / 100.0) * (unique_styles / 10.0))
        return complexity

    def _detect_document_type(self, paragraphs: List) -> str:
        """Detect the type of document based on content"""
        text_sample = " ".join(p.text for p in paragraphs[:10]).lower()
        
        if "policy" in text_sample or "procedure" in text_sample:
            return "policy_manual"
        elif "training" in text_sample or "manual" in text_sample:
            return "training_manual"
        elif "report" in text_sample:
            return "report"
        else:
            return "general_document"

    def _generate_episode_id(self, user_id: str, document_path: str) -> str:
        """Generate a unique episode ID"""
        timestamp = datetime.now().isoformat()
        content = f"{user_id}_{document_path}_{timestamp}"
        return hashlib.md5(content.encode()).hexdigest()[:12]

    def _get_document_id(self, document_path: str) -> str:
        """Generate a document ID based on path and content"""
        return Path(document_path).stem

    def _calculate_episode_summary(self) -> Dict[str, Any]:
        """Calculate summary metrics for the episode"""
        if not self.current_episode:
            return {}
        
        total_actions = len(self.current_episode.actions)
        avg_confidence = sum(a.confidence for a in self.current_episode.actions) / max(1, total_actions)
        avg_quality = sum(r.quality_score for r in self.current_episode.rewards) / max(1, len(self.current_episode.rewards))
        
        method_distribution = Counter(a.method_used for a in self.current_episode.actions)
        
        return {
            "episode_id": self.current_episode.episode_id,
            "total_actions": total_actions,
            "total_rewards": len(self.current_episode.rewards),
            "average_confidence": avg_confidence,
            "average_quality_score": avg_quality,
            "method_distribution": dict(method_distribution),
            "processing_time": self.current_episode.total_processing_time,
            "final_quality": self.current_episode.final_quality_score
        }

    def _save_episode(self, episode: Episode) -> None:
        """Save complete episode to storage"""
        with open(self.episodes_file, "a") as f:
            f.write(json.dumps(asdict(episode)) + "\n")

    def _save_state(self, state: DocumentState) -> None:
        """Save state to storage"""
        with open(self.states_file, "a") as f:
            f.write(json.dumps(asdict(state)) + "\n")

    def _save_action(self, action: FormattingAction) -> None:
        """Save action to storage"""
        with open(self.actions_file, "a") as f:
            f.write(json.dumps(asdict(action)) + "\n")

    def _save_reward(self, reward: Reward) -> None:
        """Save reward to storage"""
        with open(self.rewards_file, "a") as f:
            f.write(json.dumps(asdict(reward)) + "\n")

    def _update_quality_metrics(self, episode: Episode) -> None:
        """Update overall quality metrics"""
        self.quality_metrics["total_episodes"] += 1
        self.quality_metrics["total_actions"] += len(episode.actions)
        
        if episode.final_quality_score:
            current_avg = self.quality_metrics["average_quality_score"]
            total_episodes = self.quality_metrics["total_episodes"]
            new_avg = (current_avg * (total_episodes - 1) + episode.final_quality_score) / total_episodes
            self.quality_metrics["average_quality_score"] = new_avg
        
        # Update method performance
        method = episode.method_used
        if method not in self.quality_metrics["method_performance"]:
            self.quality_metrics["method_performance"][method] = {
                "episodes": 0,
                "avg_quality": 0.0,
                "avg_time": 0.0
            }
        
        method_stats = self.quality_metrics["method_performance"][method]
        method_stats["episodes"] += 1
        
        if episode.final_quality_score:
            current_avg = method_stats["avg_quality"]
            episodes = method_stats["episodes"]
            method_stats["avg_quality"] = (current_avg * (episodes - 1) + episode.final_quality_score) / episodes
        
        current_time_avg = method_stats["avg_time"]
        episodes = method_stats["episodes"]
        method_stats["avg_time"] = (current_time_avg * (episodes - 1) + episode.total_processing_time) / episodes
        
        # Save updated metrics
        with open(self.metadata_file, "w") as f:
            json.dump(self.quality_metrics, f, indent=2)

    def _export_trajectory_format(self) -> List[Dict[str, Any]]:
        """Export data in trajectory format for RL training"""
        trajectories = []
        
        if not self.episodes_file.exists():
            return trajectories
        
        with open(self.episodes_file, "r") as f:
            for line in f:
                episode = json.loads(line)
                
                # Create trajectory in S-A-R format
                trajectory = {
                    "episode_id": episode["episode_id"],
                    "states": episode["states"],
                    "actions": episode["actions"],
                    "rewards": episode["rewards"],
                    "metadata": episode["metadata"]
                }
                trajectories.append(trajectory)
        
        return trajectories

    def _export_state_action_format(self) -> List[Dict[str, Any]]:
        """Export data as state-action pairs"""
        pairs = []
        
        trajectories = self._export_trajectory_format()
        for trajectory in trajectories:
            states = trajectory["states"]
            actions = trajectory["actions"]
            rewards = trajectory["rewards"]
            
            # Create state-action-reward tuples
            for i in range(min(len(states) - 1, len(actions), len(rewards))):
                pairs.append({
                    "state": states[i],
                    "action": actions[i],
                    "reward": rewards[i],
                    "next_state": states[i + 1] if i + 1 < len(states) else None,
                    "episode_id": trajectory["episode_id"]
                })
        
        return pairs

    def _export_episodes_format(self) -> List[Dict[str, Any]]:
        """Export complete episodes"""
        episodes = []
        
        if self.episodes_file.exists():
            with open(self.episodes_file, "r") as f:
                for line in f:
                    episodes.append(json.loads(line))
        
        return episodes

    def _log_event(self, event_type: str, data: Dict[str, Any]) -> None:
        """Log events for debugging and monitoring"""
        log_entry = {
            "timestamp": datetime.now().isoformat(),
            "event_type": event_type,
            "data": data
        }
        
        log_file = self.data_dir / "events.log"
        with open(log_file, "a") as f:
            f.write(json.dumps(log_entry) + "\n")

# Usage example and integration helpers
def integrate_with_multi_stage_formatter():
    """Example integration with the multi-stage formatter"""
    
    from multi_stage_formatter import MultiStageFormatter
    
    class RLEnabledMultiStageFormatter(MultiStageFormatter):
        def __init__(self, rl_collector: RLDataCollector = None):
            super().__init__()
            self.rl_collector = rl_collector or RLDataCollector()
        
        def format_document_with_rl(self, input_path: str, output_path: str, 
                                   user_id: str = "system", known_path: str = None):
            """Format document while collecting RL training data"""
            
            # Start RL episode
            episode_id = self.rl_collector.start_episode(
                user_id=user_id,
                document_path=input_path,
                method_used="multi_stage",
                user_preferences={"known_target": known_path}
            )
            
            try:
                # Run original formatting with RL data collection
                start_time = time.time()
                processed, filtered, style_counts = self.format_document(input_path, output_path)
                processing_time = time.time() - start_time
                
                # Record overall action
                self.rl_collector.record_action(
                    action_type="full_document_formatting",
                    parameters={
                        "processed_paragraphs": processed,
                        "filtered_paragraphs": filtered,
                        "style_distribution": dict(style_counts)
                    },
                    confidence=0.8,
                    method_used="multi_stage"
                )
                
                # Calculate quality score
                quality_score = self._calculate_quality_score(output_path, known_path)
                
                self.rl_collector.record_reward(
                    quality_score=quality_score,
                    efficiency_metric=1.0 / (processing_time + 0.1),  # Efficiency bonus
                    compliance_score=0.9  # Based on filtering success
                )
                
                # End episode
                summary = self.rl_collector.end_episode(
                    output_document_path=output_path,
                    final_quality_score=quality_score
                )
                
                return processed, filtered, style_counts, summary
                
            except Exception as e:
                # Record failure
                self.rl_collector.record_reward(
                    quality_score=0.0,
                    error_reduction=-1.0,
                    manual_corrections_needed=1
                )
                
                self.rl_collector.end_episode(
                    output_document_path="",
                    final_quality_score=0.0
                )
                raise
        
        def _calculate_quality_score(self, output_path: str, known_path: str = None) -> float:
            """Calculate quality score for the formatted document"""
            if not known_path or not Path(known_path).exists():
                return 0.7  # Default score when no target available
            
            # Compare with known good format (simplified)
            try:
                output_doc = Document(output_path)
                known_doc = Document(known_path)
                
                output_paras = len([p for p in output_doc.paragraphs if p.text.strip()])
                known_paras = len([p for p in known_doc.paragraphs if p.text.strip()])
                
                # Simple paragraph count similarity
                similarity = 1.0 - abs(output_paras - known_paras) / max(output_paras, known_paras, 1)
                return max(0.0, min(1.0, similarity))
                
            except Exception:
                return 0.5

if __name__ == "__main__":
    # Demo usage
    collector = RLDataCollector()
    
    # Start episode
    episode_id = collector.start_episode(
        user_id="demo_user",
        document_path="/path/to/document.docx",
        method_used="multi_stage"
    )
    
    # Record some actions
    collector.record_action(
        action_type="filter_navigation",
        parameters={"filtered_count": 50},
        confidence=0.9,
        method_used="rule_based"
    )
    
    collector.record_reward(
        quality_score=0.85,
        compliance_score=0.9,
        efficiency_metric=0.8
    )
    
    # End episode
    summary = collector.end_episode(
        output_document_path="/path/to/output.docx",
        final_quality_score=0.85
    )
    
    print("Episode summary:", summary)
    
    # Export training data
    training_data = collector.get_training_data("trajectory")
    print(f"Collected {len(training_data)} trajectories")