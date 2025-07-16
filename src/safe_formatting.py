"""
Safe Formatting Engine
Applies only visual formatting changes with comprehensive validation.
"""

import re
import logging
from typing import Dict, List, Optional, Any
from dataclasses import dataclass
from docx import Document
from bs4 import BeautifulSoup

from .content_preservation import ContentSafetyLevel, ContentPreservationEngine


@dataclass
class SafeFormattingRule:
    """Rules for safe visual formatting."""
    element_type: str  # h1, h2, p, table, etc.
    css_properties: Dict[str, str]
    conditions: List[str]
    safety_level: ContentSafetyLevel


@dataclass
class FormattingChange:
    """Record of a single formatting change."""
    element_id: str
    change_type: str  # font, spacing, margin, alignment
    old_value: str
    new_value: str
    css_property: str
    rationale: str


@dataclass
class ParagraphData:
    """Detailed paragraph information."""
    index: int
    text: str
    style_name: str
    element_type: str
    formatting: Dict[str, Any]
    safety_level: ContentSafetyLevel


@dataclass
class TableData:
    """Table structure information."""
    index: int
    rows: List[List[Dict[str, Any]]]
    formatting: Dict[str, Any]


class SafeFormattingEngine:
    """Applies only visual formatting changes with comprehensive validation."""
    
    def __init__(self, learned_style_guide_path: Optional[str] = None):
        self.allowed_css_properties = {
            'visual_only': [
                'font-family', 'font-size', 'font-weight', 'font-style',
                'color', 'background-color', 'text-align', 'line-height',
                'margin', 'padding', 'border', 'border-collapse',
                'width', 'height', 'display', 'float', 'clear'
            ],
            'layout_only': [
                'margin-top', 'margin-bottom', 'margin-left', 'margin-right',
                'padding-top', 'padding-bottom', 'padding-left', 'padding-right',
                'text-indent', 'vertical-align', 'page-break-before', 'page-break-after'
            ]
        }
        
        self.prohibited_properties = [
            'content',  # Could change text
            'counter-increment', 'counter-reset',  # Could change numbering
            'quotes',  # Could change punctuation
            'text-transform'  # Could change content case
        ]
        
        self.default_style_guide = self._create_default_style_guide()
        self.learned_style_guide = None
        self.preservation_engine = ContentPreservationEngine()
        self.logger = logging.getLogger(__name__)
        
        # Load learned style guide if provided
        if learned_style_guide_path:
            self.load_learned_styles(learned_style_guide_path)
    
    def _create_default_style_guide(self) -> Dict[str, SafeFormattingRule]:
        """Create default corporate style guide."""
        return {
            'h1': SafeFormattingRule(
                element_type='h1',
                css_properties={
                    'font-family': 'Arial, sans-serif',
                    'font-size': '18px',
                    'font-weight': 'bold',
                    'color': '#000080',
                    'margin-bottom': '12px',
                    'margin-top': '16px'
                },
                conditions=['is_main_heading'],
                safety_level=ContentSafetyLevel.SAFE_TO_MODIFY
            ),
            'h2': SafeFormattingRule(
                element_type='h2', 
                css_properties={
                    'font-family': 'Arial, sans-serif',
                    'font-size': '14px',
                    'font-weight': 'bold',
                    'color': '#333333',
                    'margin-bottom': '8px',
                    'margin-top': '12px'
                },
                conditions=['is_section_heading'],
                safety_level=ContentSafetyLevel.SAFE_TO_MODIFY
            ),
            'h3': SafeFormattingRule(
                element_type='h3',
                css_properties={
                    'font-family': 'Arial, sans-serif',
                    'font-size': '13px',
                    'font-weight': 'bold',
                    'color': '#666666',
                    'margin-bottom': '6px',
                    'margin-top': '10px'
                },
                conditions=['is_subsection_heading'],
                safety_level=ContentSafetyLevel.SAFE_TO_MODIFY
            ),
            'p': SafeFormattingRule(
                element_type='p',
                css_properties={
                    'font-family': 'Arial, sans-serif',
                    'font-size': '12px',
                    'line-height': '1.15',
                    'margin-bottom': '6px'
                },
                conditions=['is_body_text'],
                safety_level=ContentSafetyLevel.SAFE_TO_MODIFY
            ),
            'table': SafeFormattingRule(
                element_type='table',
                css_properties={
                    'border-collapse': 'collapse',
                    'width': '100%',
                    'font-family': 'Arial, sans-serif',
                    'font-size': '12px',
                    'margin-bottom': '12px'
                },
                conditions=['is_data_table'],
                safety_level=ContentSafetyLevel.SAFE_TO_MODIFY
            ),
            'td': SafeFormattingRule(
                element_type='td',
                css_properties={
                    'border': '1px solid #ddd',
                    'padding': '6px',
                    'text-align': 'left'
                },
                conditions=['is_table_cell'],
                safety_level=ContentSafetyLevel.SAFE_TO_MODIFY
            ),
            'th': SafeFormattingRule(
                element_type='th',
                css_properties={
                    'border': '1px solid #ddd',
                    'padding': '6px',
                    'background-color': '#f2f2f2',
                    'font-weight': 'bold',
                    'text-align': 'left'
                },
                conditions=['is_table_header'],
                safety_level=ContentSafetyLevel.SAFE_TO_MODIFY
            )
        }
    
    def extract_content_safely(self, document_path: str) -> Dict[str, Any]:
        """Extract content while preserving original structure."""
        try:
            doc = Document(document_path)
            
            extracted_content = {
                'paragraphs': [],
                'tables': [],
                'metadata': {
                    'title': '',
                    'author': '',
                    'created': None,
                    'document_path': document_path
                },
                'structure': {
                    'heading_hierarchy': [],
                    'table_positions': [],
                    'list_structures': []
                }
            }
            
            # Extract paragraphs with formatting metadata
            heading_levels = {}
            
            for i, paragraph in enumerate(doc.paragraphs):
                if not paragraph.text.strip():
                    continue
                    
                element_type = self._classify_paragraph_type(paragraph)
                formatting = self._extract_paragraph_formatting(paragraph)
                safety_level = self.preservation_engine.assess_paragraph_safety(paragraph.text)
                
                # Track heading hierarchy
                if element_type.startswith('h'):
                    level = int(element_type[1])
                    heading_levels[level] = heading_levels.get(level, 0) + 1
                    extracted_content['structure']['heading_hierarchy'].append({
                        'level': level,
                        'text': paragraph.text,
                        'index': i
                    })
                
                para_data = ParagraphData(
                    index=i,
                    text=paragraph.text,
                    style_name=paragraph.style.name,
                    element_type=element_type,
                    formatting=formatting,
                    safety_level=safety_level
                )
                extracted_content['paragraphs'].append(para_data)
            
            # Extract tables
            for i, table in enumerate(doc.tables):
                table_data = self._extract_table_data(table, i)
                extracted_content['tables'].append(table_data)
                extracted_content['structure']['table_positions'].append(i)
            
            # Extract document metadata
            core_props = doc.core_properties
            extracted_content['metadata'].update({
                'title': core_props.title or '',
                'author': core_props.author or '',
                'created': core_props.created
            })
            
            self.logger.info(f"Extracted content: {len(extracted_content['paragraphs'])} paragraphs, {len(extracted_content['tables'])} tables")
            
            return extracted_content
            
        except Exception as e:
            self.logger.error(f"Failed to extract content from {document_path}: {e}")
            raise
    
    def generate_safe_html(self, content: Dict[str, Any], 
                          style_guide: Dict[str, SafeFormattingRule] = None) -> str:
        """Generate HTML with safe formatting applied."""
        if style_guide is None:
            style_guide = self.default_style_guide
        
        try:
            html_parts = ['<!DOCTYPE html>', '<html lang="en">', '<head>']
            html_parts.append('<meta charset="UTF-8">')
            html_parts.append('<meta name="viewport" content="width=device-width, initial-scale=1.0">')
            
            # Add title
            title = content['metadata'].get('title', 'Formatted Document')
            html_parts.append(f'<title>{self._escape_html(title)}</title>')
            
            # Generate CSS
            css = self._generate_safe_css(style_guide)
            html_parts.append(f'<style>{css}</style>')
            html_parts.append('</head><body>')
            
            # Add document title if present
            if title:
                html_parts.append(f'<h1 class="document-title">{self._escape_html(title)}</h1>')
            
            changes_made = []
            
            # Process paragraphs
            for para in content['paragraphs']:
                if para.safety_level == ContentSafetyLevel.CONTENT_CRITICAL:
                    # Preserve original formatting for critical content
                    html_parts.append(f'<p class="preserved" data-safety="critical">{self._escape_html(para.text)}</p>')
                    self.logger.debug(f"Preserved critical content: {para.text[:50]}...")
                elif para.safety_level == ContentSafetyLevel.REQUIRES_REVIEW:
                    # Mark for review but apply basic formatting
                    html_parts.append(f'<{para.element_type} class="requires-review" data-safety="review">{self._escape_html(para.text)}</{para.element_type}>')
                else:
                    # Safe to format
                    element_type = para.element_type
                    if element_type in style_guide:
                        rule = style_guide[element_type]
                        changes = self._apply_formatting_rule(para, rule)
                        changes_made.extend(changes)
                    
                    html_parts.append(f'<{element_type} data-safety="safe">{self._escape_html(para.text)}</{element_type}>')
            
            # Process tables
            for table in content['tables']:
                html_parts.append('<table>')
                for row_idx, row in enumerate(table.rows):
                    html_parts.append('<tr>')
                    for cell in row:
                        safety_level = self.preservation_engine.assess_paragraph_safety(cell['text'])
                        
                        if safety_level == ContentSafetyLevel.CONTENT_CRITICAL:
                            cell_class = 'preserved'
                            data_safety = 'critical'
                        elif safety_level == ContentSafetyLevel.REQUIRES_REVIEW:
                            cell_class = 'requires-review'
                            data_safety = 'review'
                        else:
                            cell_class = 'formatted'
                            data_safety = 'safe'
                        
                        # Use th for first row (headers), td for others
                        cell_tag = 'th' if row_idx == 0 else 'td'
                        html_parts.append(f'<{cell_tag} class="{cell_class}" data-safety="{data_safety}">{self._escape_html(cell["text"])}</{cell_tag}>')
                    
                    html_parts.append('</tr>')
                html_parts.append('</table>')
            
            # Add metadata footer
            html_parts.append('<div class="document-metadata">')
            html_parts.append(f'<p><small>Processed: {content["metadata"].get("document_path", "Unknown")}</small></p>')
            if content['metadata'].get('author'):
                html_parts.append(f'<p><small>Author: {self._escape_html(content["metadata"]["author"])}</small></p>')
            html_parts.append('</div>')
            
            html_parts.append('</body></html>')
            
            final_html = '\n'.join(html_parts)
            
            # Log changes made
            self.logger.info(f"Generated HTML with {len(changes_made)} formatting changes")
            for change in changes_made[:5]:  # Log first 5 changes
                self.logger.debug(f"Applied change: {change.element_id} - {change.change_type}")
            
            return final_html
            
        except Exception as e:
            self.logger.error(f"Failed to generate HTML: {e}")
            raise
    
    def _classify_paragraph_type(self, paragraph) -> str:
        """Classify paragraph type for appropriate formatting."""
        style_name = paragraph.style.name.lower()
        
        if 'heading 1' in style_name or 'title' in style_name:
            return 'h1'
        elif 'heading 2' in style_name:
            return 'h2'  
        elif 'heading 3' in style_name:
            return 'h3'
        elif 'heading 4' in style_name:
            return 'h4'
        elif 'heading 5' in style_name:
            return 'h5'
        elif 'heading 6' in style_name:
            return 'h6'
        elif 'list' in style_name or 'bullet' in style_name:
            return 'li'
        else:
            return 'p'
    
    def _extract_paragraph_formatting(self, paragraph) -> Dict[str, Any]:
        """Extract current formatting for comparison."""
        formatting = {}
        
        try:
            if paragraph.runs:
                first_run = paragraph.runs[0]
                formatting['font_name'] = first_run.font.name or 'Default'
                formatting['font_size'] = first_run.font.size.pt if first_run.font.size else 12
                formatting['bold'] = bool(first_run.font.bold)
                formatting['italic'] = bool(first_run.font.italic)
            else:
                formatting.update({
                    'font_name': 'Default',
                    'font_size': 12,
                    'bold': False,
                    'italic': False
                })
            
            formatting['alignment'] = str(paragraph.alignment) if paragraph.alignment else 'left'
            
            # Extract spacing information
            pf = paragraph.paragraph_format
            formatting['space_before'] = pf.space_before.pt if pf.space_before else 0
            formatting['space_after'] = pf.space_after.pt if pf.space_after else 0
            formatting['line_spacing'] = pf.line_spacing if pf.line_spacing else 1.0
            
        except Exception as e:
            self.logger.warning(f"Error extracting paragraph formatting: {e}")
            # Provide defaults
            formatting = {
                'font_name': 'Default',
                'font_size': 12,
                'bold': False,
                'italic': False,
                'alignment': 'left',
                'space_before': 0,
                'space_after': 0,
                'line_spacing': 1.0
            }
        
        return formatting
    
    def _extract_table_data(self, table, table_index: int) -> TableData:
        """Extract table data and structure."""
        rows = []
        
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cell_data = {
                    'text': cell.text,
                    'safety_level': self.preservation_engine.assess_paragraph_safety(cell.text)
                }
                row_data.append(cell_data)
            rows.append(row_data)
        
        formatting = {
            'border_style': 'solid',
            'border_width': '1px',
            'cell_padding': '6px'
        }
        
        return TableData(
            index=table_index,
            rows=rows,
            formatting=formatting
        )
    
    def _generate_safe_css(self, style_guide: Dict[str, SafeFormattingRule]) -> str:
        """Generate CSS that only affects visual presentation."""
        css_rules = []
        
        # Reset and base styles
        css_rules.append("""
body {
    font-family: Arial, sans-serif;
    line-height: 1.6;
    margin: 40px;
    max-width: 800px;
    color: #333;
}

.document-title {
    border-bottom: 2px solid #000080;
    padding-bottom: 10px;
    margin-bottom: 20px;
}""")
        
        # Generate styles from style guide
        for element_type, rule in style_guide.items():
            properties = []
            for prop, value in rule.css_properties.items():
                # Validate property is safe
                if prop in self.allowed_css_properties['visual_only'] + self.allowed_css_properties['layout_only']:
                    properties.append(f'  {prop}: {value};')
                else:
                    self.logger.warning(f"Skipping unsafe CSS property: {prop}")
            
            if properties:
                css_rules.append(f'{element_type} {{\n' + '\n'.join(properties) + '\n}')
        
        # Add safety-specific styling
        css_rules.append("""
.preserved {
    /* Original formatting preserved for safety-critical content */
    background-color: #fff9c4 !important;
    border-left: 4px solid #f0ad4e !important;
    padding: 8px !important;
    margin: 4px 0 !important;
    font-family: inherit !important;
    font-size: inherit !important;
}

.requires-review {
    /* Content that requires human review */
    background-color: #e6f3ff !important;
    border-left: 4px solid #007bff !important;
    padding: 4px !important;
}

.preserved::before {
    content: "⚠️ PRESERVED: ";
    font-weight: bold;
    color: #8a6d3b;
}

.requires-review::before {
    content: "👁️ REVIEW: ";
    font-weight: bold;
    color: #0056b3;
}

.document-metadata {
    margin-top: 40px;
    padding-top: 20px;
    border-top: 1px solid #ddd;
    color: #666;
    font-size: 11px;
}

/* Print styles */
@media print {
    .preserved, .requires-review {
        background-color: white !important;
        border-left: 2px solid #000 !important;
    }
    
    .preserved::before, .requires-review::before {
        content: "";
    }
}""")
        
        return '\n\n'.join(css_rules)
    
    def _apply_formatting_rule(self, paragraph: ParagraphData, 
                             rule: SafeFormattingRule) -> List[FormattingChange]:
        """Apply formatting rule and track changes."""
        changes = []
        current_formatting = paragraph.formatting
        
        for css_prop, new_value in rule.css_properties.items():
            # Map CSS property to current formatting key
            format_key = self._map_css_to_format_key(css_prop)
            
            if format_key and format_key in current_formatting:
                old_value = str(current_formatting[format_key])
                if old_value != new_value:
                    change = FormattingChange(
                        element_id=f"paragraph_{paragraph.index}",
                        change_type=css_prop,
                        old_value=old_value,
                        new_value=new_value,
                        css_property=css_prop,
                        rationale=f"Applied {rule.element_type} style rule"
                    )
                    changes.append(change)
        
        return changes
    
    def _map_css_to_format_key(self, css_property: str) -> Optional[str]:
        """Map CSS property to internal formatting key."""
        mapping = {
            'font-family': 'font_name',
            'font-size': 'font_size',
            'font-weight': 'bold',
            'font-style': 'italic',
            'text-align': 'alignment'
        }
        return mapping.get(css_property)
    
    def _escape_html(self, text: str) -> str:
        """Escape HTML special characters."""
        if not text:
            return ""
        
        return (text
                .replace('&', '&amp;')
                .replace('<', '&lt;')
                .replace('>', '&gt;')
                .replace('"', '&quot;')
                .replace("'", '&#x27;'))
    
    def validate_css_safety(self, css_content: str) -> Dict[str, Any]:
        """Validate that CSS only contains safe properties."""
        validation_result = {
            'is_safe': True,
            'violations': [],
            'warnings': []
        }
        
        # Check for prohibited properties
        for prohibited_prop in self.prohibited_properties:
            if prohibited_prop in css_content:
                validation_result['is_safe'] = False
                validation_result['violations'].append(f"Prohibited property found: {prohibited_prop}")
        
        # Check for potentially unsafe patterns
        unsafe_patterns = [
            r'content\s*:',  # Content modification
            r'javascript\s*:',  # JavaScript URLs
            r'expression\s*\(',  # IE expressions
            r'@import',  # External imports
        ]
        
        for pattern in unsafe_patterns:
            if re.search(pattern, css_content, re.IGNORECASE):
                validation_result['is_safe'] = False
                validation_result['violations'].append(f"Unsafe pattern found: {pattern}")
        
        return validation_result
    
    def load_learned_styles(self, style_guide_path: str):
        """Load learned style guide from JSON file."""
        try:
            from style_learning_engine import StyleLearningEngine
            
            learning_engine = StyleLearningEngine()
            self.learned_style_guide = learning_engine.load_learned_guide(style_guide_path)
            
            # Convert learned patterns to formatting rules
            learned_rules = learning_engine.convert_to_formatting_rules(self.learned_style_guide)
            
            # Merge with default style guide (learned rules take precedence)
            self.style_guide = {**self.default_style_guide, **learned_rules}
            
            self.logger.info(f"Loaded {len(learned_rules)} learned formatting rules from {style_guide_path}")
            
        except Exception as e:
            self.logger.error(f"Failed to load learned styles: {e}")
            self.style_guide = self.default_style_guide
    
    def get_active_style_guide(self) -> Dict[str, SafeFormattingRule]:
        """Get the currently active style guide (learned or default)."""
        if hasattr(self, 'style_guide'):
            return self.style_guide
        return self.default_style_guide
    
    def get_learned_patterns_info(self) -> Optional[Dict[str, Any]]:
        """Get information about loaded learned patterns."""
        if not self.learned_style_guide:
            return None
        
        return {
            'patterns_count': len(self.learned_style_guide.patterns),
            'default_font': self.learned_style_guide.default_font,
            'default_size': self.learned_style_guide.default_size,
            'confidence_threshold': self.learned_style_guide.confidence_threshold,
            'training_examples': self.learned_style_guide.training_examples,
            'pattern_types': [p.pattern_type for p in self.learned_style_guide.patterns]
        }


if __name__ == "__main__":
    # Example usage
    logging.basicConfig(level=logging.INFO)
    
    engine = SafeFormattingEngine()
    print("Safe Formatting Engine initialized")
    print(f"Default style guide has {len(engine.default_style_guide)} rules")
    
    # Show CSS safety validation
    test_css = """
    p { font-family: Arial; color: blue; }
    .bad { content: "malicious"; }
    """
    
    validation = engine.validate_css_safety(test_css)
    print(f"\nCSS Safety Validation: {validation}")