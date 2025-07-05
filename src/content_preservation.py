"""
Content Preservation Engine
Ensures zero content risk through comprehensive validation and fingerprinting.
"""

import hashlib
import json
import re
import logging
from typing import Dict, List, Optional, Any
from dataclasses import dataclass
from enum import Enum
from pathlib import Path
from docx import Document
from datetime import datetime


class ContentSafetyLevel(Enum):
    SAFE_TO_MODIFY = "safe"           # Visual formatting only
    CONTENT_CRITICAL = "critical"     # Must not be touched
    REQUIRES_REVIEW = "review"        # Human verification needed


@dataclass
class ContentFingerprint:
    """Immutable content signature for integrity verification."""
    full_text_hash: str
    paragraph_hashes: List[str]
    structure_hash: str
    word_count: int
    character_count: int
    numerical_values: List[str]
    timestamp: datetime


@dataclass
class ProhibitedZone:
    """Areas of document that must not be modified."""
    zone_type: str  # procedural, technical, numerical, regulatory
    start_paragraph: int
    end_paragraph: int
    content_hash: str
    safety_level: ContentSafetyLevel
    text_sample: str  # First 100 chars for debugging


class ContentPreservationEngine:
    """Ensures zero content risk through comprehensive validation."""
    
    def __init__(self):
        self.prohibited_patterns = {
            'procedural_language': [
                r'\b(step|procedure|process|method|instruction)\b',
                r'\b(shall|must|required|mandatory|prohibited)\b',
                r'\b(warning|caution|danger|notice|alert)\b',
                r'\b(follow|execute|perform|complete)\b.*\b(exactly|precisely)\b'
            ],
            'technical_terms': [
                r'\b[A-Z]{2,}\b',  # Acronyms
                r'\b\w+[-_]\w+\b',  # Technical notation
                r'\b\d+\.\d+\.\d+\b',  # Version numbers
                r'\b[a-zA-Z]+\d+[a-zA-Z]*\b',  # Product codes
                r'\b(API|SDK|HTTP|REST|JSON|XML)\b',
                r'\b(OAuth|SSL|TLS|HTTPS)\b'
            ],
            'numerical_values': [
                r'\b\d+\.\d+\b',  # Decimal numbers
                r'\$\d+(?:,\d{3})*(?:\.\d{2})?\b',  # Currency
                r'\b\d+%\b',  # Percentages
                r'\b\d{1,3}(?:,\d{3})*\b',  # Large numbers
                r'\b\d+\s*(PSI|V|A|Hz|°[CF])\b',  # Units
                r'\b\d+[-/]\d+[-/]\d{2,4}\b'  # Dates
            ],
            'regulatory_compliance': [
                r'\b(FDA|ISO|OSHA|EPA|compliance|regulation)\b',
                r'\b(standard|specification|requirement|guideline)\b',
                r'\b(approved|certified|validated|authorized)\b',
                r'\b(CE|UL|FCC|RoHS)\b.*\b(certified|compliant)\b'
            ]
        }
        self.logger = logging.getLogger(__name__)
    
    def create_content_fingerprint(self, document_path: str) -> ContentFingerprint:
        """Generate comprehensive content signature."""
        try:
            doc = Document(document_path)
            
            # Extract all text content
            full_text = '\n'.join(paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip())
            
            # Create content hashes
            full_text_hash = hashlib.sha256(full_text.encode('utf-8')).hexdigest()
            paragraph_hashes = [
                hashlib.sha256(p.text.encode('utf-8')).hexdigest() 
                for p in doc.paragraphs if p.text.strip()
            ]
            
            # Structure fingerprint
            structure_data = {
                'paragraph_count': len([p for p in doc.paragraphs if p.text.strip()]),
                'table_count': len(doc.tables),
                'heading_count': len([p for p in doc.paragraphs if p.style.name.startswith('Heading')])
            }
            structure_hash = hashlib.sha256(json.dumps(structure_data, sort_keys=True).encode('utf-8')).hexdigest()
            
            # Extract numerical values for protection
            numerical_values = []
            for pattern_list in self.prohibited_patterns['numerical_values']:
                matches = re.findall(pattern_list, full_text)
                numerical_values.extend(matches)
            
            # Remove duplicates while preserving order
            numerical_values = list(dict.fromkeys(numerical_values))
            
            fingerprint = ContentFingerprint(
                full_text_hash=full_text_hash,
                paragraph_hashes=paragraph_hashes,
                structure_hash=structure_hash,
                word_count=len(full_text.split()),
                character_count=len(full_text),
                numerical_values=numerical_values,
                timestamp=datetime.now()
            )
            
            self.logger.info(f"Created fingerprint for {document_path}: {len(paragraph_hashes)} paragraphs, {len(numerical_values)} numerical values")
            return fingerprint
            
        except Exception as e:
            self.logger.error(f"Failed to create fingerprint for {document_path}: {e}")
            raise
    
    def identify_prohibited_zones(self, document_path: str) -> List[ProhibitedZone]:
        """Identify areas that must not be modified."""
        try:
            doc = Document(document_path)
            prohibited_zones = []
            
            for i, paragraph in enumerate(doc.paragraphs):
                if not paragraph.text.strip():
                    continue
                    
                # Check each prohibition category
                for zone_type, patterns in self.prohibited_patterns.items():
                    for pattern in patterns:
                        if re.search(pattern, paragraph.text, re.IGNORECASE):
                            zone = ProhibitedZone(
                                zone_type=zone_type,
                                start_paragraph=i,
                                end_paragraph=i,  # Single paragraph for now
                                content_hash=hashlib.sha256(paragraph.text.encode('utf-8')).hexdigest(),
                                safety_level=ContentSafetyLevel.CONTENT_CRITICAL,
                                text_sample=paragraph.text[:100]
                            )
                            prohibited_zones.append(zone)
                            self.logger.debug(f"Found {zone_type} in paragraph {i}: {paragraph.text[:50]}...")
                            break  # Only mark once per paragraph
            
            self.logger.info(f"Identified {len(prohibited_zones)} prohibited zones in {document_path}")
            
            # Log zone summary
            zone_summary = {}
            for zone in prohibited_zones:
                zone_summary[zone.zone_type] = zone_summary.get(zone.zone_type, 0) + 1
            self.logger.info(f"Zone breakdown: {zone_summary}")
            
            return prohibited_zones
            
        except Exception as e:
            self.logger.error(f"Failed to identify prohibited zones in {document_path}: {e}")
            raise
    
    def validate_content_integrity(self, original_fingerprint: ContentFingerprint, 
                                 processed_content: str) -> Dict[str, Any]:
        """Verify content has not been modified."""
        try:
            processed_hash = hashlib.sha256(processed_content.encode('utf-8')).hexdigest()
            
            integrity_checks = {
                'content_hash_match': processed_hash == original_fingerprint.full_text_hash,
                'word_count_match': len(processed_content.split()) == original_fingerprint.word_count,
                'character_count_match': len(processed_content) == original_fingerprint.character_count,
                'numerical_values_preserved': True,
                'missing_values': []
            }
            
            # Check numerical values preservation
            for value in original_fingerprint.numerical_values:
                if value not in processed_content:
                    integrity_checks['numerical_values_preserved'] = False
                    integrity_checks['missing_values'].append(value)
                    self.logger.error(f"Numerical value lost: {value}")
            
            is_valid = (
                integrity_checks['content_hash_match'] and
                integrity_checks['word_count_match'] and 
                integrity_checks['character_count_match'] and
                integrity_checks['numerical_values_preserved']
            )
            
            integrity_checks['overall_valid'] = is_valid
            
            if not is_valid:
                self.logger.error(f"Content integrity check failed: {integrity_checks}")
            else:
                self.logger.info("Content integrity validation passed")
            
            return integrity_checks
            
        except Exception as e:
            self.logger.error(f"Content integrity validation failed: {e}")
            return {
                'overall_valid': False,
                'error': str(e)
            }
    
    def assess_paragraph_safety(self, text: str) -> ContentSafetyLevel:
        """Assess if a paragraph is safe to format."""
        if not text.strip():
            return ContentSafetyLevel.SAFE_TO_MODIFY
            
        # Check for prohibited patterns
        for category, patterns in self.prohibited_patterns.items():
            for pattern in patterns:
                if re.search(pattern, text, re.IGNORECASE):
                    return ContentSafetyLevel.CONTENT_CRITICAL
        
        # Check for any numerical content (requires review)
        if re.search(r'\d+', text):
            return ContentSafetyLevel.REQUIRES_REVIEW
        
        return ContentSafetyLevel.SAFE_TO_MODIFY
    
    def get_safety_report(self, document_path: str) -> Dict[str, Any]:
        """Generate comprehensive safety assessment report."""
        try:
            doc = Document(document_path)
            
            safety_counts = {
                ContentSafetyLevel.SAFE_TO_MODIFY: 0,
                ContentSafetyLevel.REQUIRES_REVIEW: 0,
                ContentSafetyLevel.CONTENT_CRITICAL: 0
            }
            
            total_paragraphs = 0
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    safety_level = self.assess_paragraph_safety(paragraph.text)
                    safety_counts[safety_level] += 1
                    total_paragraphs += 1
            
            prohibited_zones = self.identify_prohibited_zones(document_path)
            
            report = {
                'document_path': document_path,
                'total_paragraphs': total_paragraphs,
                'safety_breakdown': {
                    'safe_to_modify': safety_counts[ContentSafetyLevel.SAFE_TO_MODIFY],
                    'requires_review': safety_counts[ContentSafetyLevel.REQUIRES_REVIEW],
                    'content_critical': safety_counts[ContentSafetyLevel.CONTENT_CRITICAL]
                },
                'prohibited_zones_count': len(prohibited_zones),
                'zone_types': list(set(zone.zone_type for zone in prohibited_zones)),
                'automation_feasibility': {
                    'safe_percentage': (safety_counts[ContentSafetyLevel.SAFE_TO_MODIFY] / total_paragraphs * 100) if total_paragraphs > 0 else 0,
                    'recommended_approach': self._get_recommended_approach(safety_counts, total_paragraphs)
                }
            }
            
            self.logger.info(f"Safety report generated: {report['automation_feasibility']['safe_percentage']:.1f}% safe to modify")
            
            return report
            
        except Exception as e:
            self.logger.error(f"Failed to generate safety report: {e}")
            raise
    
    def _get_recommended_approach(self, safety_counts: Dict[ContentSafetyLevel, int], total_paragraphs: int) -> str:
        """Determine recommended processing approach based on safety analysis."""
        if total_paragraphs == 0:
            return "No content to process"
        
        safe_percentage = safety_counts[ContentSafetyLevel.SAFE_TO_MODIFY] / total_paragraphs * 100
        critical_percentage = safety_counts[ContentSafetyLevel.CONTENT_CRITICAL] / total_paragraphs * 100
        
        if critical_percentage > 50:
            return "Manual formatting recommended - high safety-critical content"
        elif safe_percentage > 70:
            return "Automated formatting recommended with human review"
        elif safe_percentage > 40:
            return "Semi-automated formatting with extensive human review"
        else:
            return "Manual formatting recommended - limited automation benefits"


def create_example_fingerprint():
    """Create an example fingerprint for testing."""
    engine = ContentPreservationEngine()
    
    # This would normally use a real document
    example_fingerprint = ContentFingerprint(
        full_text_hash="example_hash_12345",
        paragraph_hashes=["hash1", "hash2", "hash3"],
        structure_hash="structure_hash_67890",
        word_count=150,
        character_count=800,
        numerical_values=["$1,000.00", "45 PSI", "15%"],
        timestamp=datetime.now()
    )
    
    return example_fingerprint


if __name__ == "__main__":
    # Example usage
    logging.basicConfig(level=logging.INFO)
    
    engine = ContentPreservationEngine()
    print("Content Preservation Engine initialized")
    print("Prohibited pattern categories:", list(engine.prohibited_patterns.keys()))
    
    # Test safety assessment
    test_texts = [
        "This is a general paragraph about our company.",
        "Step 1: Turn off the main power switch immediately.",
        "API version 2.1.3 with OAuth authentication.",
        "Total cost is $15,000.00 with 15% discount.",
        "The meeting is scheduled for next Tuesday."
    ]
    
    print("\nSafety Assessment Examples:")
    for text in test_texts:
        safety_level = engine.assess_paragraph_safety(text)
        print(f"'{text[:40]}...' -> {safety_level.value}")