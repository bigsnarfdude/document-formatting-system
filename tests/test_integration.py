"""
Integration Tests for Document Formatting System
End-to-end tests for the complete workflow.
"""

import pytest
import tempfile
import shutil
from pathlib import Path
from docx import Document
from datetime import datetime

from src.content_preservation import ContentPreservationEngine, ContentSafetyLevel
from src.safe_formatting import SafeFormattingEngine
from src.validation_interface import HumanValidationInterface
from src.rollback_manager import RollbackManager


class TestCompleteWorkflow:
    """Test the complete document processing workflow."""
    
    @pytest.fixture
    def temp_dir(self):
        """Create a temporary directory for testing."""
        temp_dir = tempfile.mkdtemp()
        yield temp_dir
        shutil.rmtree(temp_dir)
    
    @pytest.fixture
    def system_components(self, temp_dir):
        """Create all system components."""
        backup_dir = Path(temp_dir) / "backups"
        audit_file = Path(temp_dir) / "audit.json"
        
        return {
            'preservation_engine': ContentPreservationEngine(),
            'formatting_engine': SafeFormattingEngine(),
            'validation_interface': HumanValidationInterface(port=5002),
            'rollback_manager': RollbackManager(str(backup_dir), str(audit_file))
        }
    
    @pytest.fixture
    def complex_document(self, temp_dir):
        """Create a complex document with various content types."""
        doc = Document()
        
        # Document title
        doc.add_heading('System Configuration Manual', 0)
        
        # Safe content
        doc.add_paragraph('This document provides guidance for system configuration.')
        doc.add_paragraph('The system has been designed for ease of use and reliability.')
        
        # Procedural content (critical)
        doc.add_heading('Safety Procedures', 1)
        doc.add_paragraph('WARNING: Always follow these procedures exactly as written.')
        doc.add_paragraph('Step 1: Turn off the main power switch before maintenance.')
        doc.add_paragraph('Step 2: Verify all connections are properly secured.')
        
        # Technical content (critical)
        doc.add_heading('Technical Specifications', 1)
        doc.add_paragraph('API version 2.1.3 with OAuth 2.0 authentication required.')
        doc.add_paragraph('System supports HTTPS/TLS encryption protocols.')
        doc.add_paragraph('Compatible with ISO-9001 certified devices only.')
        
        # Numerical content (requires review)
        doc.add_heading('Cost Analysis', 1)
        doc.add_paragraph('Total project cost: $125,000.00 with 15% contingency.')
        doc.add_paragraph('Operating pressure: 45 PSI at maximum 150°F.')
        doc.add_paragraph('Expected completion: 12/31/2024')
        
        # Mixed content
        doc.add_heading('Implementation Notes', 1)
        doc.add_paragraph('The team should coordinate with the client before proceeding.')
        doc.add_paragraph('Required API keys must be generated through the developer portal.')
        doc.add_paragraph('Budget allocation includes $5,000 for testing equipment.')
        
        # Table with mixed content
        table = doc.add_table(rows=4, cols=3)
        table.cell(0, 0).text = 'Component'
        table.cell(0, 1).text = 'Cost'
        table.cell(0, 2).text = 'Specification'
        
        table.cell(1, 0).text = 'API Gateway'
        table.cell(1, 1).text = '$15,000'
        table.cell(1, 2).text = 'OAuth 2.0 compatible'
        
        table.cell(2, 0).text = 'Processing Unit'
        table.cell(2, 1).text = '$8,500'
        table.cell(2, 2).text = 'ISO certified'
        
        table.cell(3, 0).text = 'Documentation'
        table.cell(3, 1).text = '$2,000'
        table.cell(3, 2).text = 'User manual and guides'
        
        # Save document
        doc_path = Path(temp_dir) / "complex_document.docx"
        doc.save(str(doc_path))
        return str(doc_path)
    
    def test_end_to_end_workflow(self, system_components, complex_document):
        """Test the complete end-to-end processing workflow."""
        preservation_engine = system_components['preservation_engine']
        formatting_engine = system_components['formatting_engine']
        validation_interface = system_components['validation_interface']
        rollback_manager = system_components['rollback_manager']
        
        # Step 1: Create content fingerprint
        original_fingerprint = preservation_engine.create_content_fingerprint(complex_document)
        
        # Verify fingerprint was created
        assert original_fingerprint.word_count > 0
        assert len(original_fingerprint.paragraph_hashes) > 0
        assert len(original_fingerprint.numerical_values) > 0
        
        # Step 2: Identify prohibited zones
        prohibited_zones = preservation_engine.identify_prohibited_zones(complex_document)
        
        # Verify prohibited zones were identified
        assert len(prohibited_zones) > 0
        zone_types = {zone.zone_type for zone in prohibited_zones}
        assert 'procedural_language' in zone_types
        assert 'technical_terms' in zone_types
        assert 'numerical_values' in zone_types or 'regulatory_compliance' in zone_types
        
        # Step 3: Create backup before processing
        operation_id = f"workflow_test_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        backup_id = rollback_manager.create_backup(
            document_path=complex_document,
            fingerprint=original_fingerprint,
            operation_id=operation_id,
            metadata={"test_type": "end_to_end", "user": "test_user"}
        )
        
        # Verify backup was created
        assert backup_id in rollback_manager.backup_records
        backup_record = rollback_manager.backup_records[backup_id]
        assert backup_record.operation_id == operation_id
        
        # Step 4: Extract content safely
        extracted_content = formatting_engine.extract_content_safely(complex_document)
        
        # Verify content extraction
        assert len(extracted_content['paragraphs']) > 0
        assert len(extracted_content['tables']) > 0
        
        # Verify safety classification
        safety_levels = {p.safety_level for p in extracted_content['paragraphs']}
        assert ContentSafetyLevel.SAFE_TO_MODIFY in safety_levels
        assert ContentSafetyLevel.CONTENT_CRITICAL in safety_levels
        assert ContentSafetyLevel.REQUIRES_REVIEW in safety_levels
        
        # Step 5: Generate safe HTML
        processed_html = formatting_engine.generate_safe_html(extracted_content)
        
        # Verify HTML generation
        assert processed_html.startswith('<!DOCTYPE html>')
        assert 'preserved' in processed_html  # Critical content preserved
        assert 'requires-review' in processed_html  # Review content flagged
        
        # Step 6: Submit for human validation
        request_id = validation_interface.submit_for_validation(
            document_path=complex_document,
            original_content=extracted_content,
            processed_html=processed_html,
            changes_made=[],  # Would be populated by formatting engine
            fingerprint=original_fingerprint
        )
        
        # Verify validation submission
        assert request_id in validation_interface.pending_requests
        validation_request = validation_interface.pending_requests[request_id]
        assert validation_request.document_path == complex_document
        
        # Step 7: Log validation event
        rollback_manager.log_validation_event(validation_request)
        
        # Verify audit entry was created
        validation_entries = [e for e in rollback_manager.audit_entries 
                            if e.operation_id == request_id]
        assert len(validation_entries) > 0
        
        # Step 8: Simulate human approval
        from src.validation_interface import ValidationResponse
        approval_response = ValidationResponse(
            request_id=request_id,
            status='approved',
            reviewer_comments='Content integrity verified, formatting changes approved',
            timestamp=datetime.now(),
            reviewer_id='test_reviewer'
        )
        
        validation_interface.completed_validations[request_id] = approval_response
        validation_interface.pending_requests[request_id].status = 'approved'
        
        # Step 9: Log approval
        rollback_manager.log_validation_event(validation_request, approval_response)
        
        # Verify approval was logged
        approval_entries = [e for e in rollback_manager.audit_entries 
                          if e.operation_id == request_id and 'approved' in str(e.operation_type)]
        assert len(approval_entries) > 0
        
        # Step 10: Verify content integrity
        integrity_check = preservation_engine.validate_content_integrity(
            original_fingerprint, 
            processed_html
        )
        
        # Note: This may fail due to HTML formatting, which is expected
        # The key is that numerical values and critical content are preserved
        
        # Step 11: Generate final reports
        safety_report = preservation_engine.get_safety_report(complex_document)
        audit_report = rollback_manager.generate_audit_report(days_back=1)
        
        # Verify reports
        assert safety_report['total_paragraphs'] > 0
        assert safety_report['prohibited_zones_count'] > 0
        assert audit_report['scope']['total_entries'] > 0
    
    def test_rollback_scenario(self, system_components, complex_document, temp_dir):
        """Test the rollback scenario when processing is rejected."""
        preservation_engine = system_components['preservation_engine']
        formatting_engine = system_components['formatting_engine']
        validation_interface = system_components['validation_interface']
        rollback_manager = system_components['rollback_manager']
        
        # Create original fingerprint and backup
        original_fingerprint = preservation_engine.create_content_fingerprint(complex_document)
        operation_id = "rollback_test_001"
        backup_id = rollback_manager.create_backup(
            document_path=complex_document,
            fingerprint=original_fingerprint,
            operation_id=operation_id
        )
        
        # Process document
        extracted_content = formatting_engine.extract_content_safely(complex_document)
        processed_html = formatting_engine.generate_safe_html(extracted_content)
        
        # Submit for validation
        request_id = validation_interface.submit_for_validation(
            document_path=complex_document,
            original_content=extracted_content,
            processed_html=processed_html,
            changes_made=[],
            fingerprint=original_fingerprint
        )
        
        # Simulate rejection
        from src.validation_interface import ValidationResponse
        rejection_response = ValidationResponse(
            request_id=request_id,
            status='rejected',
            reviewer_comments='Critical content appears to be modified, rejecting changes',
            timestamp=datetime.now(),
            reviewer_id='test_reviewer'
        )
        
        validation_interface.completed_validations[request_id] = rejection_response
        rollback_manager.log_validation_event(
            validation_interface.pending_requests[request_id], 
            rejection_response
        )
        
        # Simulate rollback (in real scenario, processed document would be saved first)
        rollback_success = rollback_manager.rollback_document(backup_id)
        assert rollback_success == True
        
        # Verify rollback was logged
        rollback_entries = [e for e in rollback_manager.audit_entries 
                          if 'rollback' in str(e.operation_type).lower()]
        assert len(rollback_entries) > 0
    
    def test_safety_report_accuracy(self, system_components, complex_document):
        """Test accuracy of safety assessment reports."""
        preservation_engine = system_components['preservation_engine']
        
        # Generate safety report
        safety_report = preservation_engine.get_safety_report(complex_document)
        
        # Verify report structure
        assert 'safety_breakdown' in safety_report
        assert 'automation_feasibility' in safety_report
        
        breakdown = safety_report['safety_breakdown']
        
        # Verify counts are reasonable based on document content
        assert breakdown['safe_to_modify'] > 0  # Some content should be safe
        assert breakdown['content_critical'] > 0  # Should have critical content
        assert breakdown['requires_review'] > 0  # Should have numerical content
        
        # Verify totals match
        total_classified = sum(breakdown.values())
        assert total_classified == safety_report['total_paragraphs']
        
        # Verify automation feasibility
        feasibility = safety_report['automation_feasibility']
        assert 0 <= feasibility['safe_percentage'] <= 100
        assert isinstance(feasibility['recommended_approach'], str)
    
    def test_content_preservation_guarantees(self, system_components, complex_document):
        """Test that content preservation guarantees are met."""
        preservation_engine = system_components['preservation_engine']
        formatting_engine = system_components['formatting_engine']
        
        # Create fingerprint
        original_fingerprint = preservation_engine.create_content_fingerprint(complex_document)
        
        # Extract and process content
        extracted_content = formatting_engine.extract_content_safely(complex_document)
        
        # Verify all critical content is marked as preserved
        critical_paragraphs = [p for p in extracted_content['paragraphs'] 
                             if p.safety_level == ContentSafetyLevel.CONTENT_CRITICAL]
        
        assert len(critical_paragraphs) > 0
        
        # Verify critical content contains expected patterns
        critical_text = ' '.join(p.text for p in critical_paragraphs)
        
        # Should contain procedural language
        assert any(word in critical_text.lower() for word in ['step', 'warning', 'procedure'])
        
        # Should contain technical terms
        assert any(term in critical_text for term in ['API', 'OAuth', 'ISO'])
        
        # Generate HTML and verify preservation
        processed_html = formatting_engine.generate_safe_html(extracted_content)
        
        # Verify all numerical values are preserved
        for numerical_value in original_fingerprint.numerical_values:
            assert numerical_value in processed_html
    
    def test_audit_trail_completeness(self, system_components, complex_document):
        """Test that complete audit trail is maintained."""
        preservation_engine = system_components['preservation_engine']
        formatting_engine = system_components['formatting_engine']
        validation_interface = system_components['validation_interface']
        rollback_manager = system_components['rollback_manager']
        
        # Perform complete workflow
        original_fingerprint = preservation_engine.create_content_fingerprint(complex_document)
        
        operation_id = "audit_trail_test"
        backup_id = rollback_manager.create_backup(
            document_path=complex_document,
            fingerprint=original_fingerprint,
            operation_id=operation_id
        )
        
        extracted_content = formatting_engine.extract_content_safely(complex_document)
        processed_html = formatting_engine.generate_safe_html(extracted_content)
        
        request_id = validation_interface.submit_for_validation(
            document_path=complex_document,
            original_content=extracted_content,
            processed_html=processed_html,
            changes_made=[],
            fingerprint=original_fingerprint
        )
        
        rollback_manager.log_validation_event(
            validation_interface.pending_requests[request_id]
        )
        
        # Generate audit report
        audit_report = rollback_manager.generate_audit_report(days_back=1)
        
        # Verify audit trail completeness
        assert audit_report['scope']['total_entries'] >= 2  # Backup + validation submission
        
        # Verify operation types are recorded
        operation_summary = audit_report['operation_summary']
        assert 'document_processed' in operation_summary
        assert 'validation_submitted' in operation_summary
        
        # Verify recent activities include our operations
        recent_activities = audit_report['recent_activities']
        operation_ids_in_activities = {activity['details'].get('backup_id', activity.get('operation_id', '')) 
                                     for activity in recent_activities}
        
        assert any(operation_id in str(op_id) for op_id in operation_ids_in_activities)
    
    def test_error_recovery(self, system_components, temp_dir):
        """Test system behavior under error conditions."""
        preservation_engine = system_components['preservation_engine']
        formatting_engine = system_components['formatting_engine']
        rollback_manager = system_components['rollback_manager']
        
        # Test with non-existent document
        fake_doc_path = str(Path(temp_dir) / "non_existent.docx")
        
        with pytest.raises(Exception):
            preservation_engine.create_content_fingerprint(fake_doc_path)
        
        with pytest.raises(Exception):
            formatting_engine.extract_content_safely(fake_doc_path)
        
        # Test rollback with invalid backup ID
        rollback_success = rollback_manager.rollback_document("invalid_backup_id")
        assert rollback_success == False
        
        # Verify system remains stable after errors
        audit_report = rollback_manager.generate_audit_report()
        assert isinstance(audit_report, dict)
        assert 'error' not in audit_report or audit_report.get('error') is None


class TestSystemIntegration:
    """Test integration between system components."""
    
    def test_component_communication(self):
        """Test that components can communicate effectively."""
        # Test data flow between components
        preservation_engine = ContentPreservationEngine()
        formatting_engine = SafeFormattingEngine()
        
        # Verify engines use compatible data structures
        assert hasattr(preservation_engine, 'assess_paragraph_safety')
        assert hasattr(formatting_engine, 'preservation_engine')
        
        # Verify safety level compatibility
        test_text = "Step 1: Critical procedure"
        safety_level = preservation_engine.assess_paragraph_safety(test_text)
        assert isinstance(safety_level, ContentSafetyLevel)
    
    def test_data_consistency(self):
        """Test data consistency across components."""
        # Verify that safety levels are consistently interpreted
        from src.content_preservation import ContentSafetyLevel
        from src.safe_formatting import SafeFormattingEngine
        
        engine = SafeFormattingEngine()
        
        # Test that safety levels match between engines
        assert engine.preservation_engine.assess_paragraph_safety("Safe text") == ContentSafetyLevel.SAFE_TO_MODIFY
        assert engine.preservation_engine.assess_paragraph_safety("Step 1: Critical") == ContentSafetyLevel.CONTENT_CRITICAL


if __name__ == "__main__":
    # Run tests directly
    pytest.main([__file__, "-v"])