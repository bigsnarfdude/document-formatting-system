"""
Test Safe Formatting Engine
Tests for HTML generation with visual-only formatting.
"""

import pytest
import tempfile
from pathlib import Path
from docx import Document
from bs4 import BeautifulSoup

from src.safe_formatting import (
    SafeFormattingEngine,
    SafeFormattingRule,
    FormattingChange,
    ParagraphData,
    TableData
)
from src.content_preservation import ContentSafetyLevel


class TestSafeFormattingEngine:
    """Test the safe formatting engine."""
    
    @pytest.fixture
    def engine(self):
        """Create a safe formatting engine."""
        return SafeFormattingEngine()
    
    @pytest.fixture
    def sample_document(self):
        """Create a sample Word document for testing."""
        doc = Document()
        
        # Add content with different safety levels
        doc.add_heading('Main Title', 0)
        doc.add_heading('Section Header', 1)
        doc.add_paragraph('This is safe body text that can be formatted.')
        doc.add_paragraph('WARNING: This contains procedural language.')
        doc.add_paragraph('Cost analysis shows $15,000 total with 15% discount.')
        doc.add_paragraph('API version 2.1.3 requires OAuth authentication.')
        doc.add_paragraph('More safe content for formatting tests.')
        
        # Add a table
        table = doc.add_table(rows=2, cols=3)
        table.cell(0, 0).text = 'Item'
        table.cell(0, 1).text = 'Cost'
        table.cell(0, 2).text = 'Notes'
        table.cell(1, 0).text = 'Widget'
        table.cell(1, 1).text = '$500.00'
        table.cell(1, 2).text = 'API-enabled device'
        
        # Save to temporary file
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            doc.save(tmp.name)
            return tmp.name
    
    def test_content_extraction(self, engine, sample_document):
        """Test extraction of content with safety classification."""
        content = engine.extract_content_safely(sample_document)
        
        # Verify structure
        assert 'paragraphs' in content
        assert 'tables' in content
        assert 'metadata' in content
        assert 'structure' in content
        
        # Verify paragraphs were extracted
        assert len(content['paragraphs']) > 0
        
        # Verify paragraph structure
        for para in content['paragraphs']:
            assert isinstance(para, ParagraphData)
            assert hasattr(para, 'text')
            assert hasattr(para, 'safety_level')
            assert hasattr(para, 'element_type')
            assert hasattr(para, 'formatting')
            assert isinstance(para.safety_level, ContentSafetyLevel)
        
        # Verify tables were extracted
        assert len(content['tables']) > 0
        
        # Verify table structure
        for table in content['tables']:
            assert isinstance(table, TableData)
            assert hasattr(table, 'rows')
            assert hasattr(table, 'formatting')
            assert len(table.rows) > 0
    
    def test_paragraph_type_classification(self, engine):
        """Test paragraph type classification."""
        # Mock paragraph objects
        class MockStyle:
            def __init__(self, name):
                self.name = name
        
        class MockParagraph:
            def __init__(self, style_name):
                self.style = MockStyle(style_name)
        
        test_cases = [
            ('Heading 1', 'h1'),
            ('heading 1', 'h1'),
            ('Title', 'h1'),
            ('Heading 2', 'h2'),
            ('Heading 3', 'h3'),
            ('Normal', 'p'),
            ('Body Text', 'p'),
            ('List Paragraph', 'li'),
            ('Bullet', 'li'),
        ]
        
        for style_name, expected_type in test_cases:
            mock_para = MockParagraph(style_name)
            actual_type = engine._classify_paragraph_type(mock_para)
            assert actual_type == expected_type, f"Style '{style_name}' should be '{expected_type}', got '{actual_type}'"
    
    def test_html_generation(self, engine, sample_document):
        """Test HTML generation from extracted content."""
        content = engine.extract_content_safely(sample_document)
        html = engine.generate_safe_html(content)
        
        # Verify HTML structure
        assert html.startswith('<!DOCTYPE html>')
        assert '<html lang="en">' in html
        assert '<head>' in html
        assert '<body>' in html
        assert '</body>' in html
        assert '</html>' in html
        
        # Parse HTML for validation
        soup = BeautifulSoup(html, 'html.parser')
        
        # Verify basic structure
        assert soup.find('title') is not None
        assert soup.find('style') is not None
        assert soup.find('body') is not None
        
        # Verify safety classes are applied
        preserved_elements = soup.find_all(class_='preserved')
        review_elements = soup.find_all(class_='requires-review')
        
        assert len(preserved_elements) > 0, "Should have preserved elements for critical content"
        assert len(review_elements) > 0, "Should have review elements for numerical content"
        
        # Verify data-safety attributes
        safe_elements = soup.find_all(attrs={'data-safety': 'safe'})
        critical_elements = soup.find_all(attrs={'data-safety': 'critical'})
        review_elements = soup.find_all(attrs={'data-safety': 'review'})
        
        assert len(safe_elements) > 0
        assert len(critical_elements) > 0
        assert len(review_elements) > 0
    
    def test_css_generation(self, engine):
        """Test CSS generation with safety validation."""
        css = engine._generate_safe_css(engine.default_style_guide)
        
        # Verify CSS contains expected elements
        assert 'body {' in css
        assert 'h1 {' in css
        assert 'h2 {' in css
        assert 'p {' in css
        assert 'table {' in css
        
        # Verify safety-specific styles
        assert '.preserved {' in css
        assert '.requires-review {' in css
        assert 'background-color: #fff9c4' in css  # Preserved highlighting
        assert 'background-color: #e6f3ff' in css  # Review highlighting
        
        # Verify no prohibited properties
        assert 'content:' not in css.lower()
        assert 'javascript:' not in css.lower()
        assert '@import' not in css.lower()
    
    def test_css_safety_validation(self, engine):
        """Test CSS safety validation."""
        # Safe CSS
        safe_css = """
        p { font-family: Arial; color: blue; margin: 10px; }
        h1 { font-size: 18px; font-weight: bold; }
        """
        validation = engine.validate_css_safety(safe_css)
        assert validation['is_safe'] == True
        assert len(validation['violations']) == 0
        
        # Unsafe CSS with prohibited properties
        unsafe_css = """
        p { font-family: Arial; }
        .bad { content: "malicious"; }
        .worse { javascript: alert('hack'); }
        """
        validation = engine.validate_css_safety(unsafe_css)
        assert validation['is_safe'] == False
        assert len(validation['violations']) > 0
        
        # Check specific violations
        violations = ' '.join(validation['violations'])
        assert 'content' in violations
        assert 'javascript' in violations
    
    def test_formatting_changes_tracking(self, engine):
        """Test tracking of formatting changes."""
        # Create mock paragraph data
        para_data = ParagraphData(
            index=0,
            text="Test paragraph",
            style_name="Normal",
            element_type="p",
            formatting={
                'font_name': 'Times New Roman',
                'font_size': 12,
                'bold': False
            },
            safety_level=ContentSafetyLevel.SAFE_TO_MODIFY
        )
        
        # Create formatting rule
        rule = SafeFormattingRule(
            element_type='p',
            css_properties={
                'font-family': 'Arial, sans-serif',
                'font-size': '13px',
                'font-weight': 'normal'
            },
            conditions=['is_body_text'],
            safety_level=ContentSafetyLevel.SAFE_TO_MODIFY
        )
        
        # Apply formatting and track changes
        changes = engine._apply_formatting_rule(para_data, rule)
        
        # Verify changes were tracked
        assert len(changes) > 0
        
        for change in changes:
            assert isinstance(change, FormattingChange)
            assert change.element_id == "paragraph_0"
            assert change.old_value != change.new_value
            assert change.rationale
    
    def test_table_processing(self, engine, sample_document):
        """Test processing of tables with safety classification."""
        content = engine.extract_content_safely(sample_document)
        html = engine.generate_safe_html(content)
        
        soup = BeautifulSoup(html, 'html.parser')
        tables = soup.find_all('table')
        
        assert len(tables) > 0, "Should have extracted and processed tables"
        
        # Verify table structure
        for table in tables:
            rows = table.find_all('tr')
            assert len(rows) > 0
            
            # Check for headers and data cells
            headers = table.find_all('th')
            data_cells = table.find_all('td')
            
            # Verify safety attributes on cells
            for cell in headers + data_cells:
                assert 'data-safety' in cell.attrs
                safety_level = cell.attrs['data-safety']
                assert safety_level in ['safe', 'review', 'critical']
    
    def test_content_preservation_in_html(self, engine, sample_document):
        """Test that content is preserved in HTML output."""
        content = engine.extract_content_safely(sample_document)
        html = engine.generate_safe_html(content)
        
        # Extract text from HTML
        soup = BeautifulSoup(html, 'html.parser')
        html_text = soup.get_text()
        
        # Extract original text
        original_text = '\n'.join(p.text for p in content['paragraphs'])
        
        # Verify all original content is present in HTML
        for para in content['paragraphs']:
            assert para.text in html_text, f"Original text '{para.text}' not found in HTML output"
        
        # Verify no content was modified
        # (This is a basic check - more comprehensive validation would use fingerprints)
        original_words = set(original_text.split())
        html_words = set(html_text.split())
        
        # All original words should be present (HTML may have additional navigation/metadata)
        missing_words = original_words - html_words
        assert len(missing_words) == 0, f"Missing words in HTML: {missing_words}"
    
    def test_default_style_guide(self, engine):
        """Test the default style guide configuration."""
        style_guide = engine.default_style_guide
        
        # Verify required elements are present
        required_elements = ['h1', 'h2', 'h3', 'p', 'table', 'td', 'th']
        for element in required_elements:
            assert element in style_guide
            
            rule = style_guide[element]
            assert isinstance(rule, SafeFormattingRule)
            assert rule.element_type == element
            assert len(rule.css_properties) > 0
            assert isinstance(rule.safety_level, ContentSafetyLevel)
        
        # Verify CSS properties are safe
        for rule in style_guide.values():
            for prop in rule.css_properties.keys():
                assert prop in engine.allowed_css_properties['visual_only'] + engine.allowed_css_properties['layout_only']
                assert prop not in engine.prohibited_properties
    
    def test_html_escape_functionality(self, engine):
        """Test HTML escaping of special characters."""
        test_strings = [
            ("Normal text", "Normal text"),
            ("Text with <tags>", "Text with &lt;tags&gt;"),
            ("Text with & ampersand", "Text with &amp; ampersand"),
            ("Text with \"quotes\"", "Text with &quot;quotes&quot;"),
            ("Text with 'apostrophe'", "Text with &#x27;apostrophe&#x27;"),
            ("", ""),
            (None, ""),
        ]
        
        for input_text, expected_output in test_strings:
            actual_output = engine._escape_html(input_text)
            assert actual_output == expected_output, f"Input: {input_text}, Expected: {expected_output}, Got: {actual_output}"
    
    def test_metadata_inclusion(self, engine, sample_document):
        """Test inclusion of document metadata in HTML."""
        content = engine.extract_content_safely(sample_document)
        html = engine.generate_safe_html(content)
        
        soup = BeautifulSoup(html, 'html.parser')
        
        # Check for metadata section
        metadata_section = soup.find(class_='document-metadata')
        assert metadata_section is not None, "Document metadata section should be present"
        
        # Verify document path is included
        assert content['metadata']['document_path'] in html or 'Unknown' in html
    
    def test_edge_cases(self, engine):
        """Test edge cases and error handling."""
        # Empty content
        empty_content = {
            'paragraphs': [],
            'tables': [],
            'metadata': {'title': '', 'author': '', 'document_path': 'test.docx'},
            'structure': {'heading_hierarchy': [], 'table_positions': [], 'list_structures': []}
        }
        
        html = engine.generate_safe_html(empty_content)
        assert html.startswith('<!DOCTYPE html>')
        assert '<title>' in html
        
        # Content with only safe paragraphs
        safe_content = {
            'paragraphs': [
                ParagraphData(0, "Safe text", "Normal", "p", {}, ContentSafetyLevel.SAFE_TO_MODIFY)
            ],
            'tables': [],
            'metadata': {'title': 'Test', 'author': '', 'document_path': 'test.docx'},
            'structure': {'heading_hierarchy': [], 'table_positions': [], 'list_structures': []}
        }
        
        html = engine.generate_safe_html(safe_content)
        soup = BeautifulSoup(html, 'html.parser')
        safe_elements = soup.find_all(attrs={'data-safety': 'safe'})
        assert len(safe_elements) > 0


class TestFormattingRuleApplication:
    """Test application of formatting rules."""
    
    def test_rule_condition_checking(self):
        """Test that formatting rules are applied based on conditions."""
        # This would test condition matching logic when implemented
        pass
    
    def test_cascading_rules(self):
        """Test that multiple formatting rules can be applied."""
        # This would test rule precedence and cascading when implemented
        pass
    
    def test_custom_style_guide(self):
        """Test using custom style guides."""
        engine = SafeFormattingEngine()
        
        # Create custom style guide
        custom_guide = {
            'p': SafeFormattingRule(
                element_type='p',
                css_properties={
                    'font-family': 'Times New Roman, serif',
                    'font-size': '14px',
                    'color': '#333333'
                },
                conditions=['custom_condition'],
                safety_level=ContentSafetyLevel.SAFE_TO_MODIFY
            )
        }
        
        # Test content with custom guide
        content = {
            'paragraphs': [
                ParagraphData(0, "Test", "Normal", "p", {'font_size': 12}, ContentSafetyLevel.SAFE_TO_MODIFY)
            ],
            'tables': [],
            'metadata': {'title': 'Test', 'author': '', 'document_path': 'test.docx'},
            'structure': {'heading_hierarchy': [], 'table_positions': [], 'list_structures': []}
        }
        
        html = engine.generate_safe_html(content, custom_guide)
        assert 'Times New Roman' in html
        assert 'font-size: 14px' in html


if __name__ == "__main__":
    # Run tests directly
    pytest.main([__file__, "-v"])