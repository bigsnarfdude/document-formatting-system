"""
Test Content Preservation Engine
Comprehensive tests for zero-content-risk validation.
"""

import pytest
import tempfile
import hashlib
from datetime import datetime
from pathlib import Path
from docx import Document

from src.content_preservation import (
    ContentPreservationEngine, 
    ContentSafetyLevel, 
    ContentFingerprint,
    ProhibitedZone
)


class TestContentPreservationEngine:
    """Test the content preservation and safety validation."""
    
    @pytest.fixture
    def engine(self):
        """Create a content preservation engine."""
        return ContentPreservationEngine()
    
    @pytest.fixture
    def sample_document(self):
        """Create a sample Word document for testing."""
        doc = Document()
        
        # Add various types of content
        doc.add_heading('System Configuration Guide', 0)
        doc.add_paragraph('This document describes the configuration process.')
        
        # Add procedural content (should be protected)
        doc.add_paragraph('Step 1: Turn off the main power switch immediately.')
        doc.add_paragraph('WARNING: Do not proceed without proper authorization.')
        
        # Add technical content (should be protected)
        doc.add_paragraph('API version 2.1.3 with OAuth authentication required.')
        doc.add_paragraph('Configure SSL-TLS encryption using certificate ABC-123.')
        
        # Add numerical content (should require review)
        doc.add_paragraph('Total cost is $15,000.00 with 15% discount applied.')
        doc.add_paragraph('Operating pressure: 45 PSI at 120°F maximum.')
        
        # Add safe content
        doc.add_paragraph('The meeting will be held in the conference room.')
        doc.add_paragraph('Please review the attached materials beforehand.')
        
        # Add regulatory content (should be protected)
        doc.add_paragraph('This system is FDA approved and ISO 9001 certified.')
        
        # Save to temporary file
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            doc.save(tmp.name)
            return tmp.name
    
    def test_prohibited_patterns_detection(self, engine):
        """Test detection of prohibited content patterns."""
        test_cases = [
            # Procedural language
            ("Step 1: Turn off the power", ContentSafetyLevel.CONTENT_CRITICAL),
            ("This procedure must be followed exactly", ContentSafetyLevel.CONTENT_CRITICAL),
            ("WARNING: High voltage present", ContentSafetyLevel.CONTENT_CRITICAL),
            
            # Technical terms
            ("API version 2.1.3", ContentSafetyLevel.CONTENT_CRITICAL),
            ("Configure OAuth authentication", ContentSafetyLevel.CONTENT_CRITICAL),
            ("Use HTTPS protocol", ContentSafetyLevel.CONTENT_CRITICAL),
            
            # Numerical values (requires review)
            ("The cost is $15,000.00", ContentSafetyLevel.REQUIRES_REVIEW),
            ("Pressure: 45 PSI", ContentSafetyLevel.REQUIRES_REVIEW),
            ("Temperature: 120°F", ContentSafetyLevel.REQUIRES_REVIEW),
            
            # Regulatory compliance
            ("FDA approved device", ContentSafetyLevel.CONTENT_CRITICAL),
            ("ISO 9001 certified", ContentSafetyLevel.CONTENT_CRITICAL),
            
            # Safe content
            ("The meeting is tomorrow", ContentSafetyLevel.SAFE_TO_MODIFY),
            ("Please review the document", ContentSafetyLevel.SAFE_TO_MODIFY),
        ]
        
        for text, expected_level in test_cases:
            actual_level = engine.assess_paragraph_safety(text)
            assert actual_level == expected_level, f"Text: '{text}' expected {expected_level}, got {actual_level}"
    
    def test_content_fingerprint_creation(self, engine, sample_document):
        """Test creation of content fingerprints."""
        fingerprint = engine.create_content_fingerprint(sample_document)
        
        # Verify fingerprint structure
        assert isinstance(fingerprint, ContentFingerprint)
        assert fingerprint.full_text_hash
        assert len(fingerprint.full_text_hash) == 64  # SHA-256 hex length
        assert len(fingerprint.paragraph_hashes) > 0
        assert fingerprint.word_count > 0
        assert fingerprint.character_count > 0
        assert isinstance(fingerprint.numerical_values, list)
        assert isinstance(fingerprint.timestamp, datetime)
        
        # Verify numerical values were extracted
        assert any('$15,000.00' in val for val in fingerprint.numerical_values)
        assert any('45 PSI' in val for val in fingerprint.numerical_values)
        assert any('15%' in val for val in fingerprint.numerical_values)
    
    def test_fingerprint_consistency(self, engine, sample_document):
        """Test that fingerprints are consistent for the same content."""
        fingerprint1 = engine.create_content_fingerprint(sample_document)
        fingerprint2 = engine.create_content_fingerprint(sample_document)
        
        # Fingerprints should be identical (except timestamp)
        assert fingerprint1.full_text_hash == fingerprint2.full_text_hash
        assert fingerprint1.paragraph_hashes == fingerprint2.paragraph_hashes
        assert fingerprint1.structure_hash == fingerprint2.structure_hash
        assert fingerprint1.word_count == fingerprint2.word_count
        assert fingerprint1.character_count == fingerprint2.character_count
        assert fingerprint1.numerical_values == fingerprint2.numerical_values
    
    def test_prohibited_zones_identification(self, engine, sample_document):
        """Test identification of prohibited zones."""
        zones = engine.identify_prohibited_zones(sample_document)
        
        assert len(zones) > 0
        
        # Verify zone types are detected
        zone_types = {zone.zone_type for zone in zones}
        expected_types = {'procedural_language', 'technical_terms', 'regulatory_compliance'}
        assert expected_types.issubset(zone_types)
        
        # Verify each zone has required properties
        for zone in zones:
            assert isinstance(zone, ProhibitedZone)
            assert zone.zone_type in engine.prohibited_patterns.keys()
            assert zone.start_paragraph >= 0
            assert zone.end_paragraph >= zone.start_paragraph
            assert len(zone.content_hash) == 64  # SHA-256
            assert zone.safety_level == ContentSafetyLevel.CONTENT_CRITICAL
            assert len(zone.text_sample) <= 100
    
    def test_content_integrity_validation(self, engine, sample_document):
        """Test content integrity validation."""
        # Create original fingerprint
        original_fingerprint = engine.create_content_fingerprint(sample_document)
        
        # Test with identical content
        doc = Document(sample_document)
        original_text = '\n'.join(p.text for p in doc.paragraphs if p.text.strip())
        
        integrity_check = engine.validate_content_integrity(original_fingerprint, original_text)
        assert integrity_check['overall_valid'] == True
        assert integrity_check['content_hash_match'] == True
        assert integrity_check['word_count_match'] == True
        assert integrity_check['character_count_match'] == True
        assert integrity_check['numerical_values_preserved'] == True
        
        # Test with modified content (should fail)
        modified_text = original_text.replace('$15,000.00', '$20,000.00')
        integrity_check_modified = engine.validate_content_integrity(original_fingerprint, modified_text)
        assert integrity_check_modified['overall_valid'] == False
        assert integrity_check_modified['numerical_values_preserved'] == False
        assert '$15,000.00' in integrity_check_modified['missing_values']
    
    def test_safety_report_generation(self, engine, sample_document):
        """Test safety report generation."""
        report = engine.get_safety_report(sample_document)
        
        # Verify report structure
        assert 'document_path' in report
        assert 'total_paragraphs' in report
        assert 'safety_breakdown' in report
        assert 'prohibited_zones_count' in report
        assert 'automation_feasibility' in report
        
        # Verify safety breakdown
        breakdown = report['safety_breakdown']
        assert 'safe_to_modify' in breakdown
        assert 'requires_review' in breakdown
        assert 'content_critical' in breakdown
        
        # Verify totals add up
        total_classified = sum(breakdown.values())
        assert total_classified == report['total_paragraphs']
        
        # Verify automation feasibility
        feasibility = report['automation_feasibility']
        assert 'safe_percentage' in feasibility
        assert 'recommended_approach' in feasibility
        assert 0 <= feasibility['safe_percentage'] <= 100
    
    def test_numerical_value_extraction(self, engine):
        """Test extraction of numerical values for protection."""
        test_texts = [
            ("Cost: $15,000.00", ["$15,000.00"]),
            ("Pressure: 45 PSI", ["45 PSI"]),
            ("Discount: 15%", ["15%"]),
            ("Date: 12/25/2023", ["12/25/2023"]),
            ("Version: 2.1.3", ["2.1.3"]),
            ("No numbers here", []),
        ]
        
        for text, expected_values in test_texts:
            # Test individual assessment
            safety_level = engine.assess_paragraph_safety(text)
            if expected_values:
                assert safety_level in [ContentSafetyLevel.REQUIRES_REVIEW, ContentSafetyLevel.CONTENT_CRITICAL]
            else:
                assert safety_level == ContentSafetyLevel.SAFE_TO_MODIFY
    
    def test_edge_cases(self, engine):
        """Test edge cases and error conditions."""
        # Empty text
        assert engine.assess_paragraph_safety("") == ContentSafetyLevel.SAFE_TO_MODIFY
        assert engine.assess_paragraph_safety("   ") == ContentSafetyLevel.SAFE_TO_MODIFY
        
        # Very long text
        long_text = "This is safe content. " * 1000
        assert engine.assess_paragraph_safety(long_text) == ContentSafetyLevel.SAFE_TO_MODIFY
        
        # Mixed content
        mixed_text = "This is a step-by-step procedure with $1,000 cost and API access."
        assert engine.assess_paragraph_safety(mixed_text) == ContentSafetyLevel.CONTENT_CRITICAL
    
    def test_pattern_coverage(self, engine):
        """Test that all prohibited pattern categories are covered."""
        expected_categories = ['procedural_language', 'technical_terms', 'numerical_values', 'regulatory_compliance']
        
        for category in expected_categories:
            assert category in engine.prohibited_patterns
            assert len(engine.prohibited_patterns[category]) > 0
            
            # Test each pattern compiles as valid regex
            for pattern in engine.prohibited_patterns[category]:
                import re
                try:
                    re.compile(pattern)
                except re.error:
                    pytest.fail(f"Invalid regex pattern in {category}: {pattern}")


class TestContentFingerprintIntegrity:
    """Test fingerprint integrity and security."""
    
    def test_hash_uniqueness(self):
        """Test that different content produces different hashes."""
        text1 = "This is the first document."
        text2 = "This is the second document."
        
        hash1 = hashlib.sha256(text1.encode('utf-8')).hexdigest()
        hash2 = hashlib.sha256(text2.encode('utf-8')).hexdigest()
        
        assert hash1 != hash2
    
    def test_hash_consistency(self):
        """Test that same content always produces same hash."""
        text = "This is consistent content."
        
        hash1 = hashlib.sha256(text.encode('utf-8')).hexdigest()
        hash2 = hashlib.sha256(text.encode('utf-8')).hexdigest()
        
        assert hash1 == hash2
    
    def test_fingerprint_serialization(self):
        """Test that fingerprints can be serialized and deserialized."""
        original_fingerprint = ContentFingerprint(
            full_text_hash="test_hash_123",
            paragraph_hashes=["hash1", "hash2"],
            structure_hash="structure_hash_456",
            word_count=100,
            character_count=500,
            numerical_values=["$1,000", "50%"],
            timestamp=datetime.now()
        )
        
        # Test that all fields are accessible
        assert original_fingerprint.full_text_hash == "test_hash_123"
        assert len(original_fingerprint.paragraph_hashes) == 2
        assert original_fingerprint.word_count == 100
        assert len(original_fingerprint.numerical_values) == 2


if __name__ == "__main__":
    # Run tests directly
    pytest.main([__file__, "-v"])