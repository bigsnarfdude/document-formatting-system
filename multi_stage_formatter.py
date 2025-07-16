#!/usr/bin/env python3
"""
Multi-Stage Formatter - 3-pass filtering for bulletproof content cleaning
Stage 1: Remove navigation/TOC elements
Stage 2: Remove headers/footers/metadata  
Stage 3: Classify and format content
"""

import re
from pathlib import Path
from docx import Document
from collections import Counter
import time

class MultiStageFormatter:
    """Multi-stage formatter with 3-pass filtering."""
    
    def __init__(self):
        self.stage1_filtered = []
        self.stage2_filtered = []
        self.stage3_classified = []
        self.debug_info = {"stage1": [], "stage2": [], "stage3": []}
    
    def stage1_navigation_filter(self, paragraphs):
        """Stage 1: Remove navigation and TOC elements."""
        print("🔍 Stage 1: Filtering navigation and TOC elements...")
        
        kept_paragraphs = []
        filtered_count = 0
        
        for para in paragraphs:
            text = para.text.strip()
            text_lower = text.lower()
            should_filter = False
            reason = ""
            
            # Empty or too short
            if len(text) < 3:
                should_filter = True
                reason = "too short"
            
            # Table of contents patterns
            elif text == "TABLE OF CONTENTS":
                should_filter = True
                reason = "table of contents"
            
            elif "table of contents" in text_lower:
                should_filter = True
                reason = "table of contents variant"
            
            elif text == "MASTER TABLE OF CONTENTS":
                should_filter = True
                reason = "master table of contents"
            
            elif "master table" in text_lower:
                should_filter = True
                reason = "master table variant"
            
            # Record and revision patterns
            elif "record of revisions" in text_lower:
                should_filter = True
                reason = "record of revisions"
            
            elif "revision highlights" in text_lower:
                should_filter = True
                reason = "revision highlights"
            
            elif "list of effective sections" in text_lower:
                should_filter = True
                reason = "list of effective sections"
            
            elif "list of effective pages" in text_lower:
                should_filter = True
                reason = "list of effective pages"
            
            # Dotted lines (TOC formatting)
            elif re.search(r'\.{5,}', text):
                should_filter = True
                reason = "dotted lines (TOC)"
            
            # Page references (A-1, REV-1, RH-1, LoES-1)
            elif re.match(r'^[A-Z]+-\d+$', text):
                should_filter = True
                reason = "page reference"
            
            # TOC entries with page numbers
            elif re.search(r'^\d+\.\d+\s+.*\.{3,}\s*\d+', text):
                should_filter = True
                reason = "TOC entry with page numbers"
            
            if should_filter:
                self.debug_info["stage1"].append(f"'{text}' ({reason})")
                filtered_count += 1
            else:
                kept_paragraphs.append(para)
        
        print(f"  Stage 1 complete: {len(kept_paragraphs)} kept, {filtered_count} filtered")
        return kept_paragraphs
    
    def stage2_metadata_filter(self, paragraphs):
        """Stage 2: Remove headers, footers, and metadata."""
        print("🔍 Stage 2: Filtering headers, footers, and metadata...")
        
        kept_paragraphs = []
        filtered_count = 0
        
        for para in paragraphs:
            text = para.text.strip()
            text_lower = text.lower()
            should_filter = False
            reason = ""
            
            # CRITICAL: Intentionally left blank - HIGHEST PRIORITY
            if text == "INTENTIONALLY LEFT BLANK":
                should_filter = True
                reason = "intentionally left blank (exact)"
            
            elif "intentionally left blank" in text_lower:
                should_filter = True
                reason = "intentionally left blank (variant)"
            
            elif "this page intentionally left blank" in text_lower:
                should_filter = True
                reason = "this page intentionally left blank"
            
            # Headers and footers
            elif re.search(r'page \d+', text_lower):
                should_filter = True
                reason = "page number"
            
            elif re.search(r'revision:', text_lower):
                should_filter = True
                reason = "revision info"
            
            elif re.search(r'report #', text_lower):
                should_filter = True
                reason = "report number"
            
            elif "oae." in text_lower:
                should_filter = True
                reason = "oae reference"
            
            elif "company name" in text_lower and len(text) < 50:
                should_filter = True
                reason = "company header"
            
            # Document titles
            elif text == "FLIGHT ATTENDANT POLICIES & PROCEDURES HANDBOOK":
                should_filter = True
                reason = "document title"
            
            elif "policies & procedures handbook" in text_lower and len(text) < 100:
                should_filter = True
                reason = "handbook title"
            
            elif "flight crew training manual" in text_lower and len(text) < 100:
                should_filter = True
                reason = "manual title"
            
            # Date patterns
            elif re.search(r'date:?\s*\d{2}-\d{2}-\d{2}', text_lower):
                should_filter = True
                reason = "date stamp"
            
            if should_filter:
                self.debug_info["stage2"].append(f"'{text}' ({reason})")
                filtered_count += 1
            else:
                kept_paragraphs.append(para)
        
        print(f"  Stage 2 complete: {len(kept_paragraphs)} kept, {filtered_count} filtered")
        return kept_paragraphs
    
    def stage3_classify_content(self, paragraphs):
        """Stage 3: Classify remaining content."""
        print("🔍 Stage 3: Classifying content...")
        
        classified_paragraphs = []
        
        for para in paragraphs:
            text = para.text.strip()
            text_lower = text.lower()
            
            # Classify the paragraph
            style = self._classify_paragraph(text)
            
            classified_paragraphs.append({
                'text': text,
                'style': style,
                'original_para': para
            })
            
            self.debug_info["stage3"].append(f"'{text[:50]}...' -> {style}")
        
        print(f"  Stage 3 complete: {len(classified_paragraphs)} paragraphs classified")
        return classified_paragraphs
    
    def _classify_paragraph(self, text):
        """Classify paragraph content."""
        text_clean = text.strip()
        text_lower = text_clean.lower()
        
        # Strong list indicators
        list_patterns = [
            r'^[•\-\*◦►○]\s+',
            r'^\d+\.\s+',
            r'^[a-z]\)\s+',
            r'^[A-Z]\)\s+',
            r'^\([a-z]\)\s+',
            r'^\([A-Z]\)\s+',
        ]
        
        for pattern in list_patterns:
            if re.match(pattern, text_clean):
                return "List Paragraph"
        
        # Imperative verbs (strong list indicators)
        imperative_starts = [
            'read,', 'understand', 'comply', 'report', 'maintain', 'avoid',
            'use', 'respect', 'promote', 'cooperate', 'ensure', 'follow',
            'wear', 'keep', 'replace', 'check', 'verify', 'submit'
        ]
        
        if any(text_lower.startswith(verb) for verb in imperative_starts):
            return "List Paragraph"
        
        # Modal verbs = often lists
        modal_phrases = ['must ', 'shall ', 'will ', 'should ', 'are required to', 'are responsible for']
        if any(phrase in text_lower for phrase in modal_phrases):
            return "List Paragraph"
        
        # Heading patterns - only for short text that doesn't end with period
        if len(text_clean) < 100 and not text_clean.endswith('.'):
            # All caps = likely heading
            if text_clean.isupper():
                if len(text_clean) < 30:
                    return "Heading 2"
                else:
                    return "Heading 1"
            
            # Title case = likely heading
            elif text_clean.istitle():
                return "Heading 3"
        
        # Special patterns
        if text_clean.startswith('NOTE:'):
            return "Normal"
        
        # Default to Body Text
        return "Body Text"
    
    def format_document(self, input_path, output_path):
        """Format document using 3-stage filtering."""
        print(f"🎯 Multi-Stage Formatter (3-Pass)")
        print("="*50)
        
        # Load document
        doc = Document(input_path)
        all_paragraphs = [p for p in doc.paragraphs if p.text.strip()]
        
        print(f"Input: {len(all_paragraphs)} paragraphs")
        
        start_time = time.time()
        
        # Stage 1: Navigation/TOC filtering
        stage1_paragraphs = self.stage1_navigation_filter(all_paragraphs)
        
        # Stage 2: Headers/footers/metadata filtering
        stage2_paragraphs = self.stage2_metadata_filter(stage1_paragraphs)
        
        # Stage 3: Content classification
        classified_content = self.stage3_classify_content(stage2_paragraphs)
        
        # Create new document
        new_doc = Document()
        style_counts = Counter()
        
        print(f"📝 Creating formatted document...")
        
        for item in classified_content:
            # Create paragraph
            new_para = new_doc.add_paragraph(item['text'])
            
            # Apply style
            try:
                new_para.style = item['style']
                style_counts[item['style']] += 1
            except:
                new_para.style = 'Body Text'
                style_counts['Body Text'] += 1
        
        # Save document
        print(f"💾 Saving: {output_path}")
        new_doc.save(output_path)
        
        total_time = time.time() - start_time
        
        # Calculate totals
        total_processed = len(classified_content)
        total_filtered = len(all_paragraphs) - total_processed
        stage1_filtered = len(all_paragraphs) - len(stage1_paragraphs)
        stage2_filtered = len(stage1_paragraphs) - len(stage2_paragraphs)
        
        print(f"\n📊 Multi-Stage Results:")
        print(f"  Original paragraphs: {len(all_paragraphs)}")
        print(f"  Stage 1 filtered: {stage1_filtered} (navigation/TOC)")
        print(f"  Stage 2 filtered: {stage2_filtered} (headers/metadata)")
        print(f"  Final processed: {total_processed}")
        print(f"  Total filtered: {total_filtered}")
        print(f"  Processing time: {total_time:.1f}s")
        
        return total_processed, total_filtered, style_counts
    
    def show_debug_info(self):
        """Show debug information from each stage."""
        print(f"\n🔍 Debug Information:")
        
        for stage, items in self.debug_info.items():
            print(f"\n{stage.upper()} filtered examples (first 5):")
            for item in items[:5]:
                print(f"  {item}")
            
            if len(items) > 5:
                print(f"  ... and {len(items) - 5} more")

def compare_results(output_path, known_path):
    """Compare results with known output."""
    print("\n📊 Results Comparison")
    print("="*40)
    
    # Load documents
    output_doc = Document(output_path)
    known_doc = Document(known_path)
    
    # Count paragraphs and styles
    output_paras = [p for p in output_doc.paragraphs if p.text.strip()]
    known_paras = [p for p in known_doc.paragraphs if p.text.strip()]
    
    output_styles = Counter(p.style.name for p in output_paras)
    known_styles = Counter(p.style.name for p in known_paras)
    
    print(f"Multi-Stage: {len(output_paras)} paragraphs, {len(output_styles)} styles")
    print(f"Known Target: {len(known_paras)} paragraphs, {len(known_styles)} styles")
    
    # Calculate accuracy
    para_diff = abs(len(output_paras) - len(known_paras))
    style_diff = abs(len(output_styles) - len(known_styles))
    
    print(f"\nAccuracy:")
    print(f"Paragraph gap: {para_diff}")
    print(f"Style gap: {style_diff}")
    
    # Check for "INTENTIONALLY LEFT BLANK"
    print(f"\n🔍 Critical Check:")
    intentional_blank_found = False
    for para in output_paras:
        if "INTENTIONALLY LEFT BLANK" in para.text:
            intentional_blank_found = True
            break
    
    if intentional_blank_found:
        print("❌ FILTERING FAILED: 'INTENTIONALLY LEFT BLANK' found in output")
    else:
        print("✅ FILTERING SUCCESS: No 'INTENTIONALLY LEFT BLANK' in output")
    
    # Assessment
    if para_diff <= 10 and not intentional_blank_found:
        print("🎯 EXCELLENT! Very close match with clean filtering")
    elif para_diff <= 20 and not intentional_blank_found:
        print("✅ VERY GOOD! Target achieved with clean filtering")
    elif not intentional_blank_found:
        print("🟡 Good filtering, paragraph count needs work")
    else:
        print("❌ Filtering failed - ignored content still present")
    
    return para_diff, style_diff

def main():
    """Main function."""
    # File paths
    input_path = "/Users/vincent/Desktop/watson/PPH_original.docx"
    output_path = "/Users/vincent/Desktop/watson/PPH_claude_multi_stage_formatted.docx"
    known_path = "/Users/vincent/Desktop/watson/PPH_formatted_final.docx"
    
    # Process document
    formatter = MultiStageFormatter()
    processed, filtered, style_counts = formatter.format_document(input_path, output_path)
    
    print(f"\n📊 Final Summary:")
    print(f"  • Processed: {processed} paragraphs")
    print(f"  • Filtered: {filtered} paragraphs")
    print(f"  • Styles: {len(style_counts)}")
    
    print(f"\nStyle distribution:")
    for style, count in style_counts.most_common():
        percentage = (count / processed) * 100 if processed > 0 else 0
        print(f"  {style}: {count} ({percentage:.1f}%)")
    
    # Show debug info
    formatter.show_debug_info()
    
    # Compare with target
    para_diff, style_diff = compare_results(output_path, known_path)
    
    print(f"\n🎯 Multi-Stage Assessment:")
    print(f"3-stage filtering: {para_diff} paragraph gap, {len(style_counts)} styles")
    print(f"Previous attempts: 5-108 paragraph gap")
    
    if para_diff <= 20:
        print("✅ SUCCESS! Multi-stage filtering achieved target")
    else:
        print("📈 Progress with systematic 3-stage approach")

if __name__ == '__main__':
    main()