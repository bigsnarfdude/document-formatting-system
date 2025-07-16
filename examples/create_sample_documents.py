#!/usr/bin/env python3
"""
Create Sample Documents for Testing
Generates various types of documents to demonstrate the formatting system.
"""

import os
from pathlib import Path
from docx import Document
from docx.shared import Inches


def create_examples_directory():
    """Create examples directory structure."""
    examples_dir = Path(__file__).parent
    sample_docs_dir = examples_dir / "sample_documents"
    sample_docs_dir.mkdir(exist_ok=True)
    return sample_docs_dir


def create_simple_document(output_dir: Path):
    """Create a simple document with mostly safe content."""
    doc = Document()
    
    # Title
    doc.add_heading('Simple Business Memo', 0)
    
    # Safe content
    doc.add_paragraph('This memo provides an update on the quarterly business review.')
    doc.add_paragraph('The team has been working diligently to improve our processes.')
    doc.add_paragraph('We appreciate your continued support and feedback.')
    
    # Add some basic formatting
    doc.add_heading('Next Steps', 1)
    doc.add_paragraph('Please review the attached materials.')
    doc.add_paragraph('Contact the team lead with any questions.')
    doc.add_paragraph('Schedule follow-up meetings as needed.')
    
    # Save document
    file_path = output_dir / "simple_business_memo.docx"
    doc.save(str(file_path))
    print(f"Created: {file_path}")
    return str(file_path)


def create_technical_document(output_dir: Path):
    """Create a document with technical content requiring preservation."""
    doc = Document()
    
    # Title
    doc.add_heading('API Integration Guide', 0)
    
    # Introduction (safe)
    doc.add_paragraph('This guide explains how to integrate with our system.')
    doc.add_paragraph('The integration process has been simplified for developers.')
    
    # Technical specifications (critical content)
    doc.add_heading('Authentication Requirements', 1)
    doc.add_paragraph('WARNING: All API calls must use OAuth 2.0 authentication.')
    doc.add_paragraph('API version 3.2.1 is required for compatibility.')
    doc.add_paragraph('HTTPS endpoints must be used exclusively.')
    
    # Procedural content (critical)
    doc.add_heading('Setup Procedure', 1)
    doc.add_paragraph('Step 1: Generate API keys through the developer portal.')
    doc.add_paragraph('Step 2: Configure OAuth 2.0 client credentials.')
    doc.add_paragraph('Step 3: Test authentication with GET /api/v3/auth/verify.')
    
    # Numerical data (requires review)
    doc.add_heading('Rate Limits and Costs', 1)
    doc.add_paragraph('API rate limit: 1,000 requests per minute.')
    doc.add_paragraph('Cost: $0.01 per API call above 10,000 monthly.')
    doc.add_paragraph('Timeout: 30 seconds maximum.')
    
    # Mixed content
    doc.add_heading('Support Information', 1)
    doc.add_paragraph('Our support team is available to help with integration.')
    doc.add_paragraph('Emergency support: +1-800-555-API (24/7 availability).')
    doc.add_paragraph('Standard response time: 4 hours during business hours.')
    
    # Save document
    file_path = output_dir / "api_integration_guide.docx"
    doc.save(str(file_path))
    print(f"Created: {file_path}")
    return str(file_path)


def create_procedural_document(output_dir: Path):
    """Create a document with safety procedures and warnings."""
    doc = Document()
    
    # Title
    doc.add_heading('Equipment Safety Manual', 0)
    
    # Introduction
    doc.add_paragraph('This manual contains critical safety information.')
    doc.add_paragraph('All personnel must read and understand these procedures.')
    
    # Safety warnings (critical content)
    doc.add_heading('DANGER - High Voltage', 1)
    doc.add_paragraph('WARNING: Electrical shock hazard present.')
    doc.add_paragraph('CAUTION: Turn off main power before servicing.')
    doc.add_paragraph('NOTICE: Only qualified personnel may perform maintenance.')
    
    # Step-by-step procedures (critical content)
    doc.add_heading('Shutdown Procedure', 1)
    doc.add_paragraph('Step 1: Alert all personnel of impending shutdown.')
    doc.add_paragraph('Step 2: Save all work and close applications.')
    doc.add_paragraph('Step 3: Turn off equipment in reverse startup order.')
    doc.add_paragraph('Step 4: Verify all systems are powered down.')
    
    # Technical specifications (critical)
    doc.add_heading('Operating Parameters', 1)
    doc.add_paragraph('Operating voltage: 480V ± 5% three-phase.')
    doc.add_paragraph('Maximum pressure: 150 PSI at 75°F ambient.')
    doc.add_paragraph('Emergency stop pressure: 175 PSI automatic.')
    
    # Regulatory compliance (critical)
    doc.add_heading('Compliance Information', 1)
    doc.add_paragraph('This equipment is OSHA compliant per regulation 29 CFR 1910.')
    doc.add_paragraph('FDA registration number: 12345678.')
    doc.add_paragraph('ISO 9001:2015 certified manufacturing process.')
    
    # Contact information (mixed)
    doc.add_heading('Emergency Contacts', 1)
    doc.add_paragraph('Emergency services: 911')
    doc.add_paragraph('Plant safety officer: extension 2500')
    doc.add_paragraph('Corporate safety hotline: 1-800-SAFETY-1')
    
    # Save document
    file_path = output_dir / "equipment_safety_manual.docx"
    doc.save(str(file_path))
    print(f"Created: {file_path}")
    return str(file_path)


def create_financial_document(output_dir: Path):
    """Create a document with financial data requiring careful review."""
    doc = Document()
    
    # Title
    doc.add_heading('Quarterly Financial Report', 0)
    
    # Executive summary (mixed content)
    doc.add_paragraph('This report summarizes our financial performance for Q3 2024.')
    doc.add_paragraph('Revenue exceeded expectations while maintaining cost controls.')
    
    # Financial data (requires review)
    doc.add_heading('Revenue Summary', 1)
    doc.add_paragraph('Total revenue: $2,450,000.00 (15% increase YoY).')
    doc.add_paragraph('Gross margin: 68.5% compared to 65.2% last quarter.')
    doc.add_paragraph('Operating expenses: $892,000.00 within budget.')
    
    # Create financial table
    doc.add_heading('Detailed Breakdown', 1)
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    
    # Table headers
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Category'
    hdr_cells[1].text = 'Q3 2024'
    hdr_cells[2].text = 'Q3 2023'
    
    # Table data (all numerical - requires review)
    data = [
        ['Revenue', '$2,450,000', '$2,130,000'],
        ['Expenses', '$892,000', '$945,000'],
        ['Net Income', '$1,558,000', '$1,185,000'],
        ['Margin %', '63.6%', '55.6%']
    ]
    
    for i, row_data in enumerate(data, 1):
        row_cells = table.rows[i].cells
        for j, cell_data in enumerate(row_data):
            row_cells[j].text = cell_data
    
    # Projections (numerical content)
    doc.add_heading('Q4 Projections', 1)
    doc.add_paragraph('Projected revenue: $2,800,000 - $3,200,000.')
    doc.add_paragraph('Target margin: 70% minimum.')
    doc.add_paragraph('Expected closing date: 12/31/2024.')
    
    # Notes (safe content)
    doc.add_heading('Notes', 1)
    doc.add_paragraph('These figures are preliminary and subject to audit.')
    doc.add_paragraph('Final audited results will be available in January.')
    doc.add_paragraph('Please direct questions to the finance team.')
    
    # Save document
    file_path = output_dir / "quarterly_financial_report.docx"
    doc.save(str(file_path))
    print(f"Created: {file_path}")
    return str(file_path)


def create_mixed_content_document(output_dir: Path):
    """Create a document with all types of content mixed together."""
    doc = Document()
    
    # Title
    doc.add_heading('System Integration Project Plan', 0)
    
    # Project overview (safe)
    doc.add_paragraph('This project will integrate multiple systems for improved efficiency.')
    doc.add_paragraph('The team includes developers, system administrators, and business analysts.')
    
    # Timeline with dates (requires review)
    doc.add_heading('Project Timeline', 1)
    doc.add_paragraph('Project start: 01/15/2024')
    doc.add_paragraph('Development phase: 8 weeks (320 hours total)')
    doc.add_paragraph('Testing phase: 3 weeks starting 03/15/2024')
    doc.add_paragraph('Go-live date: 04/01/2024')
    
    # Technical requirements (critical)
    doc.add_heading('Technical Requirements', 1)
    doc.add_paragraph('API compatibility: REST v2.1 with JSON payloads required.')
    doc.add_paragraph('Authentication: OAuth 2.0 with JWT tokens mandatory.')
    doc.add_paragraph('Database: PostgreSQL 13.x minimum version.')
    
    # Budget information (requires review)
    doc.add_heading('Budget Allocation', 1)
    doc.add_paragraph('Total budget: $185,000.00 approved.')
    doc.add_paragraph('Development: $125,000 (67.6% of budget).')
    doc.add_paragraph('Testing: $35,000 (18.9% of budget).')
    doc.add_paragraph('Contingency: $25,000 (13.5% of budget).')
    
    # Risk assessment (mixed content)
    doc.add_heading('Risk Assessment', 1)
    doc.add_paragraph('CRITICAL: Data migration must not cause service interruption.')
    doc.add_paragraph('The backup strategy includes hourly snapshots during transition.')
    doc.add_paragraph('Maximum acceptable downtime: 2 hours during maintenance window.')
    
    # Procedures (critical)
    doc.add_heading('Deployment Procedure', 1)
    doc.add_paragraph('Step 1: Backup all production databases completely.')
    doc.add_paragraph('Step 2: Deploy code to staging environment first.')
    doc.add_paragraph('Step 3: Run automated test suite and validate results.')
    doc.add_paragraph('Step 4: Schedule production deployment during low-usage period.')
    
    # Team information (safe)
    doc.add_heading('Team Contacts', 1)
    doc.add_paragraph('Project manager: Available weekdays 9 AM - 5 PM.')
    doc.add_paragraph('Lead developer: On-call during deployment window.')
    doc.add_paragraph('System administrator: 24/7 support during go-live week.')
    
    # Compliance notes (critical)
    doc.add_heading('Compliance Requirements', 1)
    doc.add_paragraph('System must maintain HIPAA compliance for patient data.')
    doc.add_paragraph('SOX controls required for financial reporting modules.')
    doc.add_paragraph('ISO 27001 security standards must be maintained.')
    
    # Save document
    file_path = output_dir / "system_integration_project_plan.docx"
    doc.save(str(file_path))
    print(f"Created: {file_path}")
    return str(file_path)


def main():
    """Create all sample documents."""
    print("Creating sample documents for testing...")
    
    # Create output directory
    output_dir = create_examples_directory()
    
    # Create sample documents
    documents = [
        create_simple_document(output_dir),
        create_technical_document(output_dir),
        create_procedural_document(output_dir),
        create_financial_document(output_dir),
        create_mixed_content_document(output_dir)
    ]
    
    print(f"\nCreated {len(documents)} sample documents in {output_dir}")
    print("\nSample document descriptions:")
    print("1. simple_business_memo.docx - Mostly safe content, good for basic testing")
    print("2. api_integration_guide.docx - Technical content with API specs and procedures")
    print("3. equipment_safety_manual.docx - Safety procedures and warnings (high critical content)")
    print("4. quarterly_financial_report.docx - Financial data requiring careful review")
    print("5. system_integration_project_plan.docx - Mixed content types for comprehensive testing")
    
    return documents


if __name__ == "__main__":
    main()